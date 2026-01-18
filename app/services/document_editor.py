"""
Document Editor
===============
Upload, edit, and generate documents (.doc, .docx, .txt, .md, .pdf).
"""

import os
import io
import re
import logging
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
from datetime import datetime
import tempfile

logger = logging.getLogger(__name__)

# Optional imports - will work in degraded mode if not available
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. DOCX support disabled.")

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    PDF_WRITE_AVAILABLE = True
except ImportError:
    PDF_WRITE_AVAILABLE = False

try:
    import PyPDF2
    PDF_READ_AVAILABLE = True
except ImportError:
    PDF_READ_AVAILABLE = False


class DocumentEditor:
    """Multi-format document editor."""
    
    SUPPORTED_FORMATS = {
        'txt': 'Plain Text',
        'md': 'Markdown',
        'docx': 'Microsoft Word',
        'pdf': 'PDF Document'
    }
    
    def __init__(self, storage_dir: Optional[Path] = None):
        self.storage_dir = storage_dir or Path.home() / ".pod_wizard" / "documents"
        self.storage_dir.mkdir(parents=True, exist_ok=True)
    
    def get_supported_formats(self) -> Dict[str, str]:
        """Get currently supported formats based on available libraries."""
        formats = {'txt': 'Plain Text', 'md': 'Markdown'}
        
        if DOCX_AVAILABLE:
            formats['docx'] = 'Microsoft Word'
        
        if PDF_READ_AVAILABLE or PDF_WRITE_AVAILABLE:
            formats['pdf'] = 'PDF Document'
        
        return formats
    
    def read_document(self, file_path: str = None, file_bytes: bytes = None, 
                     filename: str = None) -> Tuple[str, str]:
        """
        Read document content.
        
        Returns:
            Tuple of (content, format)
        """
        if file_path:
            ext = Path(file_path).suffix.lower().lstrip('.')
            with open(file_path, 'rb') as f:
                file_bytes = f.read()
        elif filename:
            ext = Path(filename).suffix.lower().lstrip('.')
        else:
            raise ValueError("Must provide file_path or filename")
        
        if ext == 'txt':
            return file_bytes.decode('utf-8'), 'txt'
        
        elif ext == 'md':
            return file_bytes.decode('utf-8'), 'md'
        
        elif ext == 'docx':
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx required for DOCX files")
            
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs]
            return '\n\n'.join(paragraphs), 'docx'
        
        elif ext == 'pdf':
            if not PDF_READ_AVAILABLE:
                raise ImportError("PyPDF2 required for reading PDF files")
            
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())
            return '\n\n'.join(text_parts), 'pdf'
        
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def create_document(self, content: str, format: str, 
                       title: str = None, metadata: Dict = None) -> bytes:
        """
        Create a document in the specified format.
        
        Returns:
            Document as bytes
        """
        if format == 'txt':
            return content.encode('utf-8')
        
        elif format == 'md':
            return content.encode('utf-8')
        
        elif format == 'docx':
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx required for DOCX files")
            
            doc = DocxDocument()
            
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check for headers
                    if para.startswith('# '):
                        doc.add_heading(para[2:], level=1)
                    elif para.startswith('## '):
                        doc.add_heading(para[3:], level=2)
                    elif para.startswith('### '):
                        doc.add_heading(para[4:], level=3)
                    else:
                        doc.add_paragraph(para)
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        
        elif format == 'pdf':
            if not PDF_WRITE_AVAILABLE:
                raise ImportError("reportlab required for creating PDF files")
            
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Set up text
            c.setFont("Helvetica", 12)
            
            if title:
                c.setFont("Helvetica-Bold", 18)
                c.drawCentredString(width / 2, height - 50, title)
                c.setFont("Helvetica", 12)
            
            # Write content
            y = height - 100 if title else height - 50
            line_height = 14
            margin = 72  # 1 inch
            
            for line in content.split('\n'):
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = height - margin
                
                # Handle long lines
                while len(line) > 80:
                    c.drawString(margin, y, line[:80])
                    line = line[80:]
                    y -= line_height
                    if y < margin:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y = height - margin
                
                c.drawString(margin, y, line)
                y -= line_height
            
            c.save()
            return buffer.getvalue()
        
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def convert_document(self, content: str, from_format: str, to_format: str) -> bytes:
        """Convert document between formats."""
        return self.create_document(content, to_format)
    
    def markdown_to_html(self, content: str) -> str:
        """Convert Markdown to HTML."""
        if MARKDOWN_AVAILABLE:
            return markdown.markdown(content, extensions=['tables', 'fenced_code'])
        else:
            # Basic conversion
            html = content
            # Headers
            html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
            html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
            html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
            # Bold and italic
            html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
            html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
            # Paragraphs
            html = re.sub(r'\n\n', r'</p><p>', html)
            html = f'<p>{html}</p>'
            return html
    
    def generate_document(self, prompt: str, document_type: str,
                         api_service = None) -> str:
        """Generate document content using AI."""
        if api_service is None:
            raise ValueError("API service required for document generation")
        
        system_prompts = {
            'blog_post': """Write a well-structured blog post with:
- Engaging title
- Introduction
- 3-5 main sections with headers
- Conclusion
- Use markdown formatting""",
            
            'product_description': """Write a compelling product description with:
- Attention-grabbing headline
- Key features and benefits
- Specifications if relevant
- Call to action""",
            
            'email': """Write a professional email with:
- Clear subject line
- Greeting
- Main message
- Call to action
- Professional closing""",
            
            'press_release': """Write a professional press release with:
- Headline
- Subheadline
- Dateline
- Lead paragraph (who, what, when, where, why)
- Body paragraphs with quotes
- Boilerplate
- Contact information placeholder""",
            
            'social_media': """Create social media content with:
- Multiple post variations
- Platform-specific formatting
- Hashtag suggestions
- Call to action""",
            
            'general': """Write well-structured content with clear organization and formatting."""
        }
        
        full_prompt = f"""{system_prompts.get(document_type, system_prompts['general'])}

Topic/Request: {prompt}

Write the complete document now:"""
        
        content = api_service.generate_text(full_prompt, max_tokens=2000)
        return content
    
    def save_document(self, content: str, filename: str, format: str = None) -> Path:
        """Save document to storage directory."""
        if format is None:
            format = Path(filename).suffix.lstrip('.') or 'txt'
        
        if not filename.endswith(f'.{format}'):
            filename = f"{filename}.{format}"
        
        file_path = self.storage_dir / filename
        doc_bytes = self.create_document(content, format)
        
        with open(file_path, 'wb') as f:
            f.write(doc_bytes)
        
        return file_path
    
    def list_documents(self) -> List[Dict]:
        """List all documents in storage."""
        documents = []
        for file in self.storage_dir.iterdir():
            if file.is_file() and file.suffix.lstrip('.') in self.SUPPORTED_FORMATS:
                documents.append({
                    'name': file.name,
                    'path': str(file),
                    'format': file.suffix.lstrip('.'),
                    'size': file.stat().st_size,
                    'modified': datetime.fromtimestamp(file.stat().st_mtime).isoformat()
                })
        return sorted(documents, key=lambda d: d['modified'], reverse=True)


def render_document_editor_ui():
    """Render document editor UI in Streamlit."""
    import streamlit as st
    
    st.markdown("### üìù Document Editor")
    
    editor = DocumentEditor()
    
    # Show supported formats
    formats = editor.get_supported_formats()
    st.caption(f"Supported formats: {', '.join(formats.values())}")
    
    # Tabs
    tab1, tab2, tab3, tab4 = st.tabs(["Create", "Edit", "Generate", "Library"])
    
    with tab1:
        st.markdown("#### Create New Document")
        
        doc_name = st.text_input("Document Name", placeholder="my_document")
        doc_format = st.selectbox("Format", list(formats.keys()), 
                                 format_func=lambda x: formats[x])
        
        content = st.text_area("Content", height=300, 
                              placeholder="Start typing your document...")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Save Document", type="primary", use_container_width=True):
                if doc_name and content:
                    try:
                        path = editor.save_document(content, doc_name, doc_format)
                        st.success(f"Saved: {path.name}")
                    except Exception as e:
                        st.error(f"Save failed: {e}")
                else:
                    st.warning("Please provide name and content")
        
        with col2:
            if content:
                try:
                    doc_bytes = editor.create_document(content, doc_format, title=doc_name)
                    filename = f"{doc_name or 'document'}.{doc_format}"
                    st.download_button(
                        "Download",
                        data=doc_bytes,
                        file_name=filename,
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Export failed: {e}")
    
    with tab2:
        st.markdown("#### Edit Existing Document")
        
        uploaded = st.file_uploader(
            "Upload Document",
            type=list(formats.keys()),
            key="doc_upload"
        )
        
        if uploaded:
            try:
                content, fmt = editor.read_document(
                    file_bytes=uploaded.read(),
                    filename=uploaded.name
                )
                
                st.markdown(f"**Format:** {formats.get(fmt, fmt)}")
                
                edited_content = st.text_area(
                    "Edit Content",
                    value=content,
                    height=400,
                    key="edit_content"
                )
                
                # Export options
                export_format = st.selectbox(
                    "Export As",
                    list(formats.keys()),
                    format_func=lambda x: formats[x],
                    key="export_format"
                )
                
                if edited_content:
                    doc_bytes = editor.create_document(edited_content, export_format)
                    base_name = Path(uploaded.name).stem
                    st.download_button(
                        f"Download as {formats[export_format]}",
                        data=doc_bytes,
                        file_name=f"{base_name}.{export_format}",
                        use_container_width=True
                    )
                    
            except Exception as e:
                st.error(f"Failed to read document: {e}")
    
    with tab3:
        st.markdown("#### Generate Document with AI")
        
        doc_type = st.selectbox("Document Type", [
            "blog_post", "product_description", "email", 
            "press_release", "social_media", "general"
        ], format_func=lambda x: x.replace('_', ' ').title())
        
        prompt = st.text_area(
            "What should the document be about?",
            placeholder="Describe the topic, key points, and any specific requirements...",
            height=100
        )
        
        if st.button("Generate Document", type="primary"):
            if prompt:
                with st.spinner("Generating..."):
                    try:
                        # Get API service
                        from app.services.platform_helpers import _get_replicate_token
                        from app.services.api_service import ReplicateAPI
                        
                        token = _get_replicate_token()
                        if token:
                            api = ReplicateAPI(api_token=token)
                            content = editor.generate_document(prompt, doc_type, api)
                            
                            st.session_state['generated_doc'] = content
                            st.success("Document generated!")
                        else:
                            st.error("API key required")
                    except Exception as e:
                        st.error(f"Generation failed: {e}")
            else:
                st.warning("Please describe what you want to generate")
        
        # Show generated content
        if 'generated_doc' in st.session_state:
            st.markdown("---")
            st.markdown("**Generated Content:**")
            
            content = st.text_area(
                "Edit generated content",
                value=st.session_state['generated_doc'],
                height=300,
                key="gen_edit"
            )
            
            col1, col2 = st.columns(2)
            with col1:
                export_fmt = st.selectbox("Format", list(formats.keys()),
                                         format_func=lambda x: formats[x],
                                         key="gen_export_fmt")
            with col2:
                if content:
                    doc_bytes = editor.create_document(content, export_fmt)
                    st.download_button(
                        "Download",
                        data=doc_bytes,
                        file_name=f"generated.{export_fmt}",
                        use_container_width=True
                    )
    
    with tab4:
        st.markdown("#### Document Library")
        
        documents = editor.list_documents()
        
        if documents:
            for doc in documents:
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    st.markdown(f"**{doc['name']}**")
                    st.caption(f"{formats.get(doc['format'], doc['format'])} ‚Ä¢ {doc['size']} bytes")
                
                with col2:
                    with open(doc['path'], 'rb') as f:
                        st.download_button(
                            "üì•",
                            data=f.read(),
                            file_name=doc['name'],
                            key=f"dl_{doc['name']}"
                        )
                
                with col3:
                    if st.button("üóëÔ∏è", key=f"del_{doc['name']}"):
                        Path(doc['path']).unlink()
                        st.rerun()
        else:
            st.info("No documents yet. Create one to get started!")
