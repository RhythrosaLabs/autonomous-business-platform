"""
OTTO ENGINE - Advanced AI Assistant with Service Integration - ENHANCED EDITION v3.0
=====================================================================================
A hyperintelligent AI assistant that can:
1. Understand complex multi-step requests with enhanced NLP
2. Execute workflows autonomously with real-time feedback
3. Sync with services (YouTube, Printify, Shopify)
4. Use browser automation for web tasks
5. Display results inline in chat
6. Cache results for faster repeated operations
7. Execute parallel tasks where dependencies allow
8. Generate ANY file type (PDF, DOCX, code, configs, etc.)
9. Call specific AI models via /model commands
10. Chain workflows together with /chain command

ðŸš€ ENHANCED FEATURES v3.0:
- NEW: AI Assistants integration (/brand, /outreach, /campaign, /design, /content)
- NEW: Knowledge Base system with image analysis and document reading
- NEW: Real-time app awareness (page tracking, state monitoring)
- Slash commands for file generation (/python, /pdf, /docx, etc.)
- Slash commands for media generation (/music, /image, /video, /speak)
- Direct AI model access (/flux, /ideogram, /musicgen, etc.)
- Workflow chaining with /chain command
- Code execution capabilities
- Web research integration

SLASH COMMANDS:
ðŸ“„ Documents: /python, /html, /js, /md, /pdf, /doc, /docx, /txt, /rtf
ðŸ“Š Data: /csv, /json, /xml, /yaml, /ini, /log, /xls, /xlsx
ðŸ’» Code: /c, /cpp, /java, /bat, /sh, /ps1
ðŸŽ¨ Media: /music, /image, /video, /speak, /sound, /3d
ðŸ¤– Models: /flux, /sdxl, /ideogram, /musicgen, /kling, etc.
ðŸ§  AI Assistants: /brand, /outreach, /campaign, /design, /content
ðŸ“š Knowledge: /learn, /remember, /recall, /forget, /knowledge
ðŸ”— Workflow: /chain <step1> | <step2> | <step3>

This is the core brain of Otto Mate.
"""

from app.tabs.abp_imports_common import (
    st, os, json, logging, asyncio, uuid, re, time, hashlib, base64, tempfile,
    Path, Dict, List, Any, Optional, Tuple, Callable, Union, datetime,
    ThreadPoolExecutor, dataclass, field, Enum, lru_cache, BytesIO, pickle,
    setup_logger
)

logger = setup_logger(__name__)

from app.tabs.abp_imports_common import (
    st, os, json, logging, asyncio, uuid, re, time, hashlib, base64, tempfile,
    Path, Dict, List, Any, Optional, Tuple, Callable, Union, datetime,
    ThreadPoolExecutor, dataclass, field, Enum, lru_cache, BytesIO, pickle,
    setup_logger
)

logger = setup_logger(__name__)

# Import all models and support classes from modularized file
from .otto_engine_models import (
    OttoKnowledgeBase, OttoAppAwareness, OttoPerformanceConfig, RequestCache,
    TaskType, TaskStep, TaskPlan, OTTO_CONFIG, get_knowledge_base
)

# Re-export for backward compatibility
__all__ = [
    'OttoKnowledgeBase', 'OttoAppAwareness', 'OttoPerformanceConfig', 
    'RequestCache', 'TaskType', 'TaskStep', 'TaskPlan', 'OTTO_CONFIG',
    'get_knowledge_base', 'OttoEngine', 'SlashCommandProcessor',
]


# ============================================================================
# NOTE: Data models are now in otto_engine_models.py
# ============================================================================
# The following classes have been moved to otto_engine_models.py:
# - OttoKnowledgeBase (knowledge and memory system)
# - OttoAppAwareness (app state tracking)
# - OttoPerformanceConfig (performance tuning)
# - RequestCache (query caching)
# - TaskType, TaskStep, TaskPlan (task workflows)
#
# These are imported and re-exported above for backward compatibility.
# Remove the duplicate class definitions below:
class OttoKnowledgeBase_Reference:
    """
    Otto's knowledge base for persistent memory and context.
    Supports image analysis, document reading, and memory management.
    """
    
    def __init__(self, storage_path: str = "otto_knowledge"):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)
        self.memory_file = self.storage_path / "memory.json"
        self.documents_dir = self.storage_path / "documents"
        self.images_dir = self.storage_path / "images"
        self.documents_dir.mkdir(exist_ok=True)
        self.images_dir.mkdir(exist_ok=True)
        
        # Load existing memory
        self.memory: Dict[str, Any] = self._load_memory()
        self.documents: Dict[str, Dict] = self.memory.get("documents", {})
        self.images: Dict[str, Dict] = self.memory.get("images", {})
        self.facts: List[Dict] = self.memory.get("facts", [])
        self.context: Dict[str, Any] = self.memory.get("context", {})
    
    def _load_memory(self) -> Dict:
        """Load memory from disk."""
        if self.memory_file.exists():
            try:
                with open(self.memory_file, 'r') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Failed to load memory: {e}")
        return {"documents": {}, "images": {}, "facts": [], "context": {}}
    
    def _save_memory(self):
        """Persist memory to disk."""
        self.memory = {
            "documents": self.documents,
            "images": self.images,
            "facts": self.facts,
            "context": self.context,
            "last_updated": datetime.now().isoformat()
        }
        try:
            with open(self.memory_file, 'w') as f:
                json.dump(self.memory, f, indent=2, default=str)
        except Exception as e:
            logger.error(f"Failed to save memory: {e}")

    def get_stats(self) -> Dict:
        """Get statistics about the knowledge base."""
        return {
            "documents": len(self.documents),
            "images": len(self.images),
            "facts": len(self.facts),
            "chunks": 0,
            "last_updated": self.memory.get("last_updated", "Never")
        }

    def reindex(self):
        """Re-index the knowledge base (placeholder)."""
        self._save_memory()
        return True
    
    def analyze_image(self, image_data: bytes, filename: str, replicate_api) -> Dict:
        """
        Analyze an uploaded image using vision AI.
        Returns description, objects detected, colors, and more.
        """
        try:
            # Save image
            image_id = str(uuid.uuid4())[:8]
            image_path = self.images_dir / f"{image_id}_{filename}"
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            # Encode for API
            image_b64 = base64.b64encode(image_data).decode()
            
            # Use vision model to analyze
            analysis_prompt = """Analyze this image in detail. Provide:
1. **Description**: What does the image show?
2. **Objects**: List main objects/elements visible
3. **Colors**: Dominant colors in the image
4. **Style**: Art style, photography type, or design style
5. **Text**: Any text visible in the image
6. **Mood**: The emotional feeling or atmosphere
7. **Use Cases**: What could this image be used for (marketing, product, social media, etc.)

Be specific and detailed."""

            # Call vision API
            analysis_result = replicate_api.analyze_image(image_b64, analysis_prompt)
            
            # Store in knowledge base
            image_record = {
                "id": image_id,
                "filename": filename,
                "path": str(image_path),
                "analysis": analysis_result,
                "uploaded_at": datetime.now().isoformat(),
                "size_bytes": len(image_data)
            }
            self.images[image_id] = image_record
            self._save_memory()
            
            return {
                "success": True,
                "image_id": image_id,
                "analysis": analysis_result,
                "message": f"âœ… Image analyzed and saved to knowledge base with ID: {image_id}"
            }
            
        except Exception as e:
            logger.error(f"Image analysis failed: {e}")
            return {"success": False, "error": str(e)}
    
    def read_document(self, file_data: bytes, filename: str, replicate_api) -> Dict:
        """
        Read and parse a document (PDF, DOCX, TXT, etc.).
        Extracts content and adds to knowledge base.
        """
        try:
            doc_id = str(uuid.uuid4())[:8]
            doc_path = self.documents_dir / f"{doc_id}_{filename}"
            with open(doc_path, 'wb') as f:
                f.write(file_data)
            
            content = ""
            file_ext = filename.lower().split('.')[-1]
            
            # Extract text based on file type
            if file_ext == 'txt':
                content = file_data.decode('utf-8', errors='ignore')
            
            elif file_ext == 'pdf':
                try:
                    import PyPDF2
                    from io import BytesIO
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_data))
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
                except ImportError:
                    # Fallback: use AI to describe
                    content = "[PDF content - PyPDF2 not installed]"
            
            elif file_ext in ['doc', 'docx']:
                try:
                    from docx import Document
                    from io import BytesIO
                    doc = Document(BytesIO(file_data))
                    for para in doc.paragraphs:
                        content += para.text + "\n"
                except ImportError:
                    content = "[DOCX content - python-docx not installed]"
            
            elif file_ext in ['json']:
                content = file_data.decode('utf-8', errors='ignore')
                try:
                    parsed = json.loads(content)
                    content = json.dumps(parsed, indent=2)
                except:
                    pass
            
            elif file_ext in ['csv']:
                content = file_data.decode('utf-8', errors='ignore')
            
            elif file_ext in ['md', 'markdown']:
                content = file_data.decode('utf-8', errors='ignore')
            
            else:
                # Try to read as text
                try:
                    content = file_data.decode('utf-8', errors='ignore')
                except:
                    content = f"[Binary file - cannot extract text]"
            
            # Summarize if content is long
            summary = content[:500] if len(content) > 500 else content
            if len(content) > 500:
                summary_prompt = f"""Summarize this document content in 3-5 bullet points:

{content[:3000]}

Provide a concise summary of the key information."""
                summary = replicate_api.generate_text(summary_prompt, max_tokens=500, temperature=0.3)
            
            # Store in knowledge base
            doc_record = {
                "id": doc_id,
                "filename": filename,
                "path": str(doc_path),
                "content": content[:50000],  # Limit stored content
                "summary": summary,
                "uploaded_at": datetime.now().isoformat(),
                "size_bytes": len(file_data),
                "file_type": file_ext
            }
            self.documents[doc_id] = doc_record
            self._save_memory()
            
            return {
                "success": True,
                "document_id": doc_id,
                "summary": summary,
                "content_length": len(content),
                "message": f"âœ… Document read and added to knowledge base with ID: {doc_id}"
            }
            
        except Exception as e:
            logger.error(f"Document reading failed: {e}")
            return {"success": False, "error": str(e)}
    
    def add_fact(self, fact: str, category: str = "general") -> Dict:
        """Add a fact or piece of information to memory."""
        fact_record = {
            "id": str(uuid.uuid4())[:8],
            "fact": fact,
            "category": category,
            "added_at": datetime.now().isoformat()
        }
        self.facts.append(fact_record)
        self._save_memory()
        return {"success": True, "message": f"âœ… Fact added to memory: {fact[:50]}..."}
    
    def recall(self, query: str, replicate_api) -> Dict:
        """Search knowledge base for relevant information."""
        results = {
            "facts": [],
            "documents": [],
            "images": []
        }
        
        query_lower = query.lower()
        
        # Search facts
        for fact in self.facts:
            if query_lower in fact.get("fact", "").lower():
                results["facts"].append(fact)
        
        # Search documents
        for doc_id, doc in self.documents.items():
            if (query_lower in doc.get("filename", "").lower() or
                query_lower in doc.get("summary", "").lower() or
                query_lower in doc.get("content", "")[:5000].lower()):
                results["documents"].append({
                    "id": doc_id,
                    "filename": doc["filename"],
                    "summary": doc.get("summary", "")[:300]
                })
        
        # Search images
        for img_id, img in self.images.items():
            analysis = img.get("analysis", "")
            if isinstance(analysis, str) and query_lower in analysis.lower():
                results["images"].append({
                    "id": img_id,
                    "filename": img["filename"],
                    "analysis": analysis[:300]
                })
        
        return {
            "success": True,
            "query": query,
            "results": results,
            "total_matches": len(results["facts"]) + len(results["documents"]) + len(results["images"])
        }
    
    def get_context_summary(self) -> str:
        """Get a summary of all knowledge for context injection."""
        summary_parts = []
        
        if self.facts:
            summary_parts.append(f"**Known Facts ({len(self.facts)}):**")
            for fact in self.facts[-5:]:  # Last 5 facts
                summary_parts.append(f"- {fact['fact'][:100]}")
        
        if self.documents:
            summary_parts.append(f"\n**Documents ({len(self.documents)}):**")
            for doc_id, doc in list(self.documents.items())[-3:]:
                summary_parts.append(f"- {doc['filename']}: {doc.get('summary', '')[:100]}")
        
        if self.images:
            summary_parts.append(f"\n**Images ({len(self.images)}):**")
            for img_id, img in list(self.images.items())[-3:]:
                summary_parts.append(f"- {img['filename']}: {str(img.get('analysis', ''))[:100]}")
        
        return "\n".join(summary_parts) if summary_parts else "No knowledge stored yet."
    
    def clear_memory(self, category: str = "all") -> Dict:
        """Clear knowledge base (all or specific category)."""
        if category == "all" or category == "facts":
            self.facts = []
        if category == "all" or category == "documents":
            self.documents = {}
        if category == "all" or category == "images":
            self.images = {}
        self._save_memory()
        return {"success": True, "message": f"âœ… Cleared {category} from knowledge base"}


# ============================================================================
# APP STATE AWARENESS - Real-time monitoring of app state
# ============================================================================
class OttoAppAwareness:
    """
    Monitors and tracks the current state of the app.
    Otto uses this to understand what's happening in the app.
    """
    
    @staticmethod
    def get_current_state() -> Dict[str, Any]:
        """Get comprehensive current app state."""
        state = {
            "timestamp": datetime.now().isoformat(),
            "current_page": st.session_state.get("current_page", "Unknown"),
            "user_activity": {}
        }
        
        # Check what page/section user is on
        if "nav_selection" in st.session_state:
            state["navigation"] = st.session_state.nav_selection
        
        # Get sidebar state
        state["sidebar_tab"] = st.session_state.get("sidebar_tab", 0)
        
        # Get recent activity
        if "otto_messages" in st.session_state:
            state["chat_messages_count"] = len(st.session_state.otto_messages)
        
        # Check for active campaigns/content
        state["campaigns"] = {
            "active_campaign": st.session_state.get("current_campaign_name", None),
            "generated_images_count": len(st.session_state.get("generated_images", [])),
            "generated_videos_count": len(st.session_state.get("generated_videos", []))
        }
        
        # Check integrations status
        state["integrations"] = {
            "printify_connected": "printify_api_key" in st.session_state or os.getenv("PRINTIFY_API_KEY"),
            "shopify_connected": "shopify_store_url" in st.session_state or os.getenv("SHOPIFY_STORE_URL"),
            "youtube_connected": st.session_state.get("youtube_authenticated", False),
            "replicate_connected": "replicate_api_key" in st.session_state or get_api_key("REPLICATE_API_TOKEN")
        }
        
        # Brand info if available
        if "brand_name" in st.session_state:
            state["brand"] = {
                "name": st.session_state.get("brand_name", ""),
                "tagline": st.session_state.get("brand_tagline", ""),
                "colors": st.session_state.get("brand_colors", []),
                "voice": st.session_state.get("brand_voice", "")
            }
        
        # Recent designs
        if "product_designs" in st.session_state:
            state["recent_designs"] = len(st.session_state.product_designs)
        
        # Shopping cart / products
        if "printify_products" in st.session_state:
            state["printify_products_count"] = len(st.session_state.printify_products)
        
        # Content pipeline
        state["content_pipeline"] = {
            "scheduled_posts": len(st.session_state.get("scheduled_posts", [])),
            "draft_content": len(st.session_state.get("draft_content", []))
        }
        
        return state
    
    @staticmethod
    def get_state_summary() -> str:
        """Get a human-readable summary of current state."""
        state = OttoAppAwareness.get_current_state()
        
        summary_parts = [
            f"ðŸ“ **Current Page:** {state.get('current_page', 'Main Dashboard')}",
            f"ðŸ”— **Integrations:** "
        ]
        
        integrations = state.get("integrations", {})
        connected = []
        if integrations.get("printify_connected"):
            connected.append("Printify")
        if integrations.get("shopify_connected"):
            connected.append("Shopify")
        if integrations.get("youtube_connected"):
            connected.append("YouTube")
        summary_parts[-1] += ", ".join(connected) if connected else "None connected"
        
        campaigns = state.get("campaigns", {})
        if campaigns.get("active_campaign"):
            summary_parts.append(f"ðŸŽ¯ **Active Campaign:** {campaigns['active_campaign']}")
        
        summary_parts.append(f"ðŸ–¼ï¸ **Generated Images:** {campaigns.get('generated_images_count', 0)}")
        summary_parts.append(f"ðŸŽ¬ **Generated Videos:** {campaigns.get('generated_videos_count', 0)}")
        
        brand = state.get("brand", {})
        if brand.get("name"):
            summary_parts.append(f"ðŸ·ï¸ **Brand:** {brand['name']}")
        
        return "\n".join(summary_parts)
    
    @staticmethod
    def track_action(action: str, details: Dict = None):
        """Track user actions for context."""
        if "otto_action_history" not in st.session_state:
            st.session_state.otto_action_history = []
        
        st.session_state.otto_action_history.append({
            "action": action,
            "details": details or {},
            "timestamp": datetime.now().isoformat()
        })
        
        # Keep only last 50 actions
        st.session_state.otto_action_history = st.session_state.otto_action_history[-50:]


# Global knowledge base instance
_knowledge_base: Optional[OttoKnowledgeBase] = None

def get_knowledge_base() -> OttoKnowledgeBase:
    """Get or create the global knowledge base."""
    global _knowledge_base
    if _knowledge_base is None:
        _knowledge_base = OttoKnowledgeBase()
    return _knowledge_base


# ============================================================================
# SLASH COMMAND REGISTRY - All supported commands
# ============================================================================
SLASH_COMMANDS = {
    # Help
    "help": {"type": "help", "desc": "Show available slash commands"},
    "commands": {"type": "help", "desc": "Show available slash commands"},
    
    # AI Assistants - NEW in v3.0
    "brand": {"type": "assistant", "assistant": "brand_builder", "desc": "AI Brand Builder - generate colors, voice, names, full kit"},
    "outreach": {"type": "assistant", "assistant": "outreach", "desc": "AI Outreach Assistant - emails, DMs, lead scoring"},
    "campaign": {"type": "assistant", "assistant": "campaign", "desc": "AI Campaign Creator - goals, A/B variants, competitor analysis"},
    "design": {"type": "assistant", "assistant": "design", "desc": "AI Product Studio - enhanced prompts, trending ideas"},
    "content": {"type": "assistant", "assistant": "content", "desc": "AI Content Generator - viral hooks, scoring, variations"},
    
    # Knowledge Base - NEW in v3.0
    "learn": {"type": "knowledge", "action": "add", "desc": "Add information to knowledge base"},
    "remember": {"type": "knowledge", "action": "add", "desc": "Remember a fact or information"},
    "recall": {"type": "knowledge", "action": "recall", "desc": "Search knowledge base for information"},
    "forget": {"type": "knowledge", "action": "clear", "desc": "Clear knowledge base"},
    "knowledge": {"type": "knowledge", "action": "status", "desc": "Show knowledge base status"},
    "status": {"type": "awareness", "action": "status", "desc": "Show current app state and context"},
    
    # Printify Integration - ENHANCED v3.0
    "printify": {"type": "printify", "action": "menu", "desc": "Printify commands menu"},
    "upload": {"type": "printify", "action": "upload", "desc": "Upload design to Printify"},
    "product": {"type": "printify", "action": "create", "desc": "Create product on Printify from design"},
    "publish": {"type": "printify", "action": "publish", "desc": "Publish product to store"},
    "mockup": {"type": "printify", "action": "mockup", "desc": "Get product mockups"},
    "blueprint": {"type": "printify", "action": "blueprint", "desc": "Search product blueprints"},
    "shop": {"type": "printify", "action": "shop", "desc": "Show shop info and products"},
    
    # Shopify Integration - ENHANCED v3.0
    "shopify": {"type": "shopify", "action": "menu", "desc": "Shopify commands menu"},
    "blog": {"type": "shopify", "action": "blog", "desc": "Create SEO blog post on Shopify"},
    "store": {"type": "shopify", "action": "store", "desc": "Show store info and stats"},
    "seo": {"type": "shopify", "action": "seo", "desc": "Generate SEO content for products"},
    "inventory": {"type": "shopify", "action": "inventory", "desc": "Check inventory levels"},
    "analytics": {"type": "shopify", "action": "analytics", "desc": "Get store analytics"},
    
    # Browser Automation - ENHANCED v3.0
    "browse": {"type": "browser", "action": "browse", "desc": "Browse URL and extract info"},
    "scrape": {"type": "browser", "action": "scrape", "desc": "Scrape data from webpage"},
    "post": {"type": "browser", "action": "post", "desc": "Post to social media platform"},
    "login": {"type": "browser", "action": "login", "desc": "Login to website (uses saved credentials)"},
    "automate": {"type": "browser", "action": "automate", "desc": "Run custom browser automation"},
    "screenshot": {"type": "browser", "action": "screenshot", "desc": "Take screenshot of webpage"},
    
    # Document Generation
    "python": {"type": "code", "ext": "py", "desc": "Generate Python script"},
    "html": {"type": "code", "ext": "html", "desc": "Generate HTML file"},
    "js": {"type": "code", "ext": "js", "desc": "Generate JavaScript file"},
    "ts": {"type": "code", "ext": "ts", "desc": "Generate TypeScript file"},
    "jsx": {"type": "code", "ext": "jsx", "desc": "Generate JSX file"},
    "tsx": {"type": "code", "ext": "tsx", "desc": "Generate TSX file"},
    "css": {"type": "code", "ext": "css", "desc": "Generate CSS stylesheet"},
    "scss": {"type": "code", "ext": "scss", "desc": "Generate SCSS stylesheet"},
    "less": {"type": "code", "ext": "less", "desc": "Generate LESS stylesheet"},
    "md": {"type": "code", "ext": "md", "desc": "Generate Markdown file"},
    "pdf": {"type": "document", "ext": "pdf", "desc": "Generate PDF document"},
    "doc": {"type": "document", "ext": "doc", "desc": "Generate DOC file"},
    "docx": {"type": "document", "ext": "docx", "desc": "Generate DOCX file"},
    "txt": {"type": "code", "ext": "txt", "desc": "Generate text file"},
    "rtf": {"type": "document", "ext": "rtf", "desc": "Generate RTF file"},
    "csv": {"type": "data", "ext": "csv", "desc": "Generate CSV file"},
    "json": {"type": "data", "ext": "json", "desc": "Generate JSON file"},
    "xml": {"type": "data", "ext": "xml", "desc": "Generate XML file"},
    "yaml": {"type": "data", "ext": "yaml", "desc": "Generate YAML file"},
    "yml": {"type": "data", "ext": "yml", "desc": "Generate YML file"},
    "toml": {"type": "data", "ext": "toml", "desc": "Generate TOML file"},
    "ini": {"type": "data", "ext": "ini", "desc": "Generate INI file"},
    "env": {"type": "data", "ext": "env", "desc": "Generate .env file"},
    "log": {"type": "code", "ext": "log", "desc": "Generate log file"},
    "xls": {"type": "spreadsheet", "ext": "xls", "desc": "Generate XLS file"},
    "xlsx": {"type": "spreadsheet", "ext": "xlsx", "desc": "Generate XLSX file"},
    
    # Code Generation
    "c": {"type": "code", "ext": "c", "desc": "Generate C source code"},
    "cpp": {"type": "code", "ext": "cpp", "desc": "Generate C++ source code"},
    "h": {"type": "code", "ext": "h", "desc": "Generate C header file"},
    "hpp": {"type": "code", "ext": "hpp", "desc": "Generate C++ header file"},
    "java": {"type": "code", "ext": "java", "desc": "Generate Java source code"},
    "kt": {"type": "code", "ext": "kt", "desc": "Generate Kotlin file"},
    "swift": {"type": "code", "ext": "swift", "desc": "Generate Swift file"},
    "go": {"type": "code", "ext": "go", "desc": "Generate Go file"},
    "rs": {"type": "code", "ext": "rs", "desc": "Generate Rust file"},
    "rb": {"type": "code", "ext": "rb", "desc": "Generate Ruby file"},
    "php": {"type": "code", "ext": "php", "desc": "Generate PHP file"},
    "pl": {"type": "code", "ext": "pl", "desc": "Generate Perl file"},
    "lua": {"type": "code", "ext": "lua", "desc": "Generate Lua file"},
    "r": {"type": "code", "ext": "r", "desc": "Generate R script"},
    "scala": {"type": "code", "ext": "scala", "desc": "Generate Scala file"},
    "dart": {"type": "code", "ext": "dart", "desc": "Generate Dart file"},
    "sql": {"type": "code", "ext": "sql", "desc": "Generate SQL file"},
    "graphql": {"type": "code", "ext": "graphql", "desc": "Generate GraphQL schema"},
    "proto": {"type": "code", "ext": "proto", "desc": "Generate Protocol Buffers"},
    "bat": {"type": "code", "ext": "bat", "desc": "Generate BAT file"},
    "sh": {"type": "code", "ext": "sh", "desc": "Generate Shell script"},
    "bash": {"type": "code", "ext": "sh", "desc": "Generate Bash script"},
    "zsh": {"type": "code", "ext": "zsh", "desc": "Generate Zsh script"},
    "ps1": {"type": "code", "ext": "ps1", "desc": "Generate PowerShell script"},
    "dockerfile": {"type": "code", "ext": "Dockerfile", "desc": "Generate Dockerfile"},
    "makefile": {"type": "code", "ext": "Makefile", "desc": "Generate Makefile"},
    "cmake": {"type": "code", "ext": "CMakeLists.txt", "desc": "Generate CMake file"},
    
    # Web & Config
    "vue": {"type": "code", "ext": "vue", "desc": "Generate Vue component"},
    "svelte": {"type": "code", "ext": "svelte", "desc": "Generate Svelte component"},
    "astro": {"type": "code", "ext": "astro", "desc": "Generate Astro component"},
    "nginx": {"type": "code", "ext": "conf", "desc": "Generate Nginx config"},
    "htaccess": {"type": "code", "ext": ".htaccess", "desc": "Generate .htaccess file"},
    "gitignore": {"type": "code", "ext": ".gitignore", "desc": "Generate .gitignore file"},
    
    # Media Generation
    "music": {"type": "media", "model": "meta/musicgen", "desc": "Generate music"},
    "image": {"type": "media", "model": "flux", "desc": "Generate image"},
    "video": {"type": "media", "model": "kling", "desc": "Generate video"},
    "sora": {"type": "media", "model": "sora", "desc": "Generate video with Sora"},
    "speak": {"type": "media", "model": "minimax/speech-02-hd", "desc": "Generate speech audio"},
    "tts": {"type": "media", "model": "minimax/speech-02-hd", "desc": "Text-to-speech"},
    "sound": {"type": "media", "model": "sound", "desc": "Generate sound effect"},
    "sfx": {"type": "media", "model": "sound", "desc": "Generate sound effect"},
    "3d": {"type": "media", "model": "tencent/hunyuan3d-2", "desc": "Generate 3D model"},
    "model3d": {"type": "media", "model": "tencent/hunyuan3d-2", "desc": "Generate 3D model"},
    
    # Task Queue & Automation - NEW v3.1
    "task": {"type": "task", "action": "create", "desc": "Create a new task in the queue"},
    "queue": {"type": "task", "action": "list", "desc": "Show task queue status"},
    "run": {"type": "task", "action": "run", "desc": "Run/execute a pending task"},
    "pause": {"type": "task", "action": "pause", "desc": "Pause a running task"},
    "cancel": {"type": "task", "action": "cancel", "desc": "Cancel a task"},
    "batch": {"type": "task", "action": "batch", "desc": "Create batch of similar tasks"},
    "schedule": {"type": "task", "action": "schedule", "desc": "Schedule a task for later"},
    
    # Workflows - NEW v3.1
    "workflow": {"type": "workflow_cmd", "action": "menu", "desc": "Workflow commands menu"},
    "workflows": {"type": "workflow_cmd", "action": "list", "desc": "List saved workflows"},
    "newflow": {"type": "workflow_cmd", "action": "create", "desc": "Create a new workflow"},
    "runflow": {"type": "workflow_cmd", "action": "run", "desc": "Run a workflow by name"},
    "automate": {"type": "workflow_cmd", "action": "auto", "desc": "AI auto-generates a workflow"},
    
    # Playground - NEW v3.1
    "playground": {"type": "playground", "action": "menu", "desc": "Playground commands menu"},
    "test": {"type": "playground", "action": "test", "desc": "Test a prompt with model"},
    "compare": {"type": "playground", "action": "compare", "desc": "Compare outputs from multiple models"},
    "preset": {"type": "playground", "action": "preset", "desc": "Use or save prompt presets"},
    "experiment": {"type": "playground", "action": "experiment", "desc": "Run experiment with variations"},
    
    # Workflow
    "chain": {"type": "workflow", "desc": "Chain multiple commands together"},
}

# Command categories for autocomplete
COMMAND_CATEGORIES = {
    "ðŸ§  AI Assistants": ["brand", "outreach", "campaign", "design", "content"],
    "ðŸ“š Knowledge": ["learn", "remember", "recall", "forget", "knowledge", "status"],
    "ðŸ“‹ Task Queue": ["task", "queue", "run", "pause", "cancel", "batch", "schedule"],
    "ðŸ”§ Workflows": ["workflow", "workflows", "newflow", "runflow", "automate"],
    "ðŸŽ® Playground": ["playground", "test", "compare", "preset", "experiment"],
    "ðŸ›ï¸ Printify": ["printify", "upload", "product", "publish", "mockup", "blueprint", "shop"],
    "ðŸª Shopify": ["shopify", "blog", "store", "seo", "inventory", "analytics"],
    "ðŸŒ Browser": ["browse", "scrape", "post", "login", "automate", "screenshot"],
    "ðŸ“„ Documents": ["pdf", "doc", "docx", "txt", "rtf", "md"],
    "ðŸ’» Code": ["python", "js", "ts", "html", "css", "java", "c", "cpp", "go", "rs", "swift", "kt", "rb", "php"],
    "ðŸ“Š Data": ["json", "csv", "xml", "yaml", "yml", "toml", "xlsx", "xls", "sql"],
    "ðŸŽ¨ Media": ["image", "video", "music", "speak", "tts", "sound", "sfx", "3d"],
    "âš™ï¸ Config": ["dockerfile", "makefile", "nginx", "gitignore", "env", "ini"],
    "ðŸŒ Web": ["html", "css", "scss", "jsx", "tsx", "vue", "svelte", "astro"],
    "ðŸ“œ Scripts": ["sh", "bash", "zsh", "ps1", "bat"],
    "ðŸ”— Chaining": ["chain", "help", "commands"],
}

# AI Model Registry - Direct model access via /modelname
AI_MODELS = {
    # 3D Models
    "hunyuan3d": "tencent/hunyuan3d-2",
    "luciddreamer": "jd7h/luciddreamer",
    "hunyuan3d2": "ndreca/hunyuan3d-2.1",
    "vggt": "vufinder/vggt-1b",
    "morphix3d": "subhash25rawat/morphix3d",
    "rodin": "hyper3d/rodin",
    
    # Music Models
    "musicgen": "meta/musicgen",
    "musicgen-finetune": "sakemin/musicgen-fine-tuner",
    "lyria": "google/lyria-2",
    "minimax-music": "minimax/music-1.5",
    "musicgen-looper": "andreasjansson/musicgen-looper",
    "flux-music": "zsxkib/flux-music",
    "stable-audio": "stability-ai/stable-audio-2.5",
    
    # Speech Models
    "minimax-speech": "minimax/speech-02-hd",
    "speech": "minimax/speech-02-hd",
    
    # Text Models
    "gpt-oss": "openai/gpt-oss-120b",
    
    # Image Models
    "flux": "prunaai/flux-fast",
    "flux-fast": "prunaai/flux-fast",
    "flux-dev": "black-forest-labs/flux-dev",
    "flux-schnell": "black-forest-labs/flux-schnell",
    "sdxl": "stability-ai/sdxl",
    "seedream": "bytedance/seedream-4",
    "imagen4": "google/imagen-4-ultra",
    "bria": "bria/image-3.2",
    "flux-kontext": "black-forest-labs/flux-kontext-pro",
    "ideogram": "ideogram-ai/ideogram-v2",
    
    # Ads & Marketing
    "flux-ads": "loolau/flux-static-ads",
    "logo-context": "subhash25rawat/logo-in-context",
    "product-ads": "pipeline-examples/ads-for-products",
    "ad-inpaint": "logerzhu/ad-inpaint",
    
    # Video Models
    "kling": "kwaivgi/kling-v2.5-turbo-pro",
    "sora": "openai/sora",
    "wan": "wavymulder/wan-2.1",
    "minimax-video": "minimax/video-01",
    "luma": "luma/ray-2-720p",
    "ltx": "lightricks/ltx-video",
    "mochi": "genmo/mochi-1-preview",
}


# ============================================================================
# SMART AUTOCOMPLETE SYSTEM
# ============================================================================
def get_command_suggestions(partial_input: str, max_suggestions: int = 8) -> List[Dict[str, str]]:
    """
    Get smart autocomplete suggestions based on partial input.
    Returns list of suggestions with command, description, and category.
    """
    suggestions = []
    partial = partial_input.lower().strip()
    
    # Remove leading slash if present
    if partial.startswith('/'):
        partial = partial[1:]
    
    if not partial:
        # Show popular commands when just / is typed
        popular = ["image", "video", "python", "pdf", "music", "json", "help", "chain"]
        for cmd in popular:
            if cmd in SLASH_COMMANDS:
                suggestions.append({
                    "command": f"/{cmd}",
                    "desc": SLASH_COMMANDS[cmd]["desc"],
                    "category": _get_command_category(cmd)
                })
        return suggestions[:max_suggestions]
    
    # Exact prefix match first (highest priority)
    for cmd, info in SLASH_COMMANDS.items():
        if cmd.startswith(partial):
            suggestions.append({
                "command": f"/{cmd}",
                "desc": info["desc"],
                "category": _get_command_category(cmd),
                "priority": 0  # Highest
            })
    
    # Also check AI models
    for model_name in AI_MODELS.keys():
        if model_name.startswith(partial):
            suggestions.append({
                "command": f"/{model_name}",
                "desc": f"AI Model: {AI_MODELS[model_name]}",
                "category": "ðŸ¤– AI Models",
                "priority": 1
            })
    
    # Fuzzy match (contains) - lower priority
    if len(suggestions) < max_suggestions:
        for cmd, info in SLASH_COMMANDS.items():
            if partial in cmd and not cmd.startswith(partial):
                suggestions.append({
                    "command": f"/{cmd}",
                    "desc": info["desc"],
                    "category": _get_command_category(cmd),
                    "priority": 2
                })
    
    # Sort by priority and return
    suggestions.sort(key=lambda x: x.get("priority", 99))
    return suggestions[:max_suggestions]


def _get_command_category(cmd: str) -> str:
    """Get the category for a command."""
    for category, commands in COMMAND_CATEGORIES.items():
        if cmd in commands:
            return category
    return "ðŸ“¦ Other"


def get_all_commands_for_help() -> Dict[str, List[Dict[str, str]]]:
    """Get all commands organized by category for help display."""
    categorized = {}
    
    for category, commands in COMMAND_CATEGORIES.items():
        categorized[category] = []
        for cmd in commands:
            if cmd in SLASH_COMMANDS:
                categorized[category].append({
                    "command": f"/{cmd}",
                    "desc": SLASH_COMMANDS[cmd]["desc"]
                })
    
    # Add AI models category
    categorized["ðŸ¤– AI Models"] = []
    for model_name, model_id in AI_MODELS.items():
        categorized["ðŸ¤– AI Models"].append({
            "command": f"/{model_name}",
            "desc": model_id
        })
    
    return categorized


# ============================================================================
# PERFORMANCE CONFIGURATION
# ============================================================================
@dataclass
class OttoPerformanceConfig:
    """Tunable performance settings for Otto Engine"""
    # Speed settings
    max_concurrent_tasks: int = 4  # Parallel task execution
    api_timeout: int = 60  # Seconds before API call timeout
    cache_ttl: int = 3600  # Cache TTL in seconds (1 hour)
    
    # Retry settings
    max_retries: int = 3
    retry_base_delay: float = 1.0
    retry_max_delay: float = 30.0
    
    # AI settings
    default_temperature: float = 0.3  # Lower = more focused
    max_tokens_analysis: int = 1500  # Increased for better analysis
    max_tokens_content: int = 1200  # For content generation
    
    # Parallel execution
    enable_parallel_steps: bool = True
    parallel_batch_size: int = 3

OTTO_CONFIG = OttoPerformanceConfig()

# ============================================================================
# REQUEST CACHE FOR REPEATED QUERIES
# ============================================================================
class RequestCache:
    """Simple in-memory cache for repeated requests."""
    def __init__(self, ttl: int = 3600):
        self._cache: Dict[str, Tuple[Any, float]] = {}
        self.ttl = ttl
    
    def _hash_key(self, request: str) -> str:
        return hashlib.md5(request.lower().strip().encode()).hexdigest()
    
    def get(self, request: str) -> Optional[Any]:
        key = self._hash_key(request)
        if key in self._cache:
            value, timestamp = self._cache[key]
            if time.time() - timestamp < self.ttl:
                logger.info(f"ðŸŽ¯ Cache hit for request")
                return value
            del self._cache[key]
        return None
    
    def set(self, request: str, value: Any):
        key = self._hash_key(request)
        self._cache[key] = (value, time.time())
    
    def clear(self):
        self._cache.clear()

# Global cache instance
_request_cache = RequestCache(ttl=OTTO_CONFIG.cache_ttl)


# ============================================================================
# SLASH COMMAND PROCESSOR - Handles all /command inputs
# ============================================================================
class SlashCommandProcessor:
    """
    Process slash commands and generate appropriate outputs.
    Supports file generation, media creation, model calls, and workflow chaining.
    """
    
    def __init__(self, replicate_api):
        self.replicate = replicate_api
        self.output_dir = Path("otto_outputs")
        self.output_dir.mkdir(exist_ok=True)
        # Also save to campaigns/slash_commands for file library access
        self.campaigns_dir = Path("campaigns/slash_commands")
        self.campaigns_dir.mkdir(parents=True, exist_ok=True)
    
    def _save_to_campaigns(self, filename: str, content: bytes = None, url: str = None) -> str:
        """Save generated content to campaigns directory for file library access."""
        try:
            filepath = self.campaigns_dir / filename
            
            if content:
                with open(filepath, 'wb') as f:
                    f.write(content)
            elif url:
                # Download from URL
                import requests
                response = requests.get(url, timeout=60)
                if response.status_code == 200:
                    with open(filepath, 'wb') as f:
                        f.write(response.content)
            
            return str(filepath)
        except Exception as e:
            logger.warning(f"Failed to save to campaigns: {e}")
            return None
    
    def parse_command(self, message: str) -> Tuple[Optional[str], str]:
        """Parse a message to extract slash command and prompt."""
        message = message.strip()
        if not message.startswith('/'):
            return None, message
        
        # Extract command and prompt
        parts = message[1:].split(' ', 1)
        command = parts[0].lower()
        prompt = parts[1] if len(parts) > 1 else ""
        
        return command, prompt
    
    def is_slash_command(self, message: str) -> bool:
        """Check if message is a slash command."""
        if not message.strip().startswith('/'):
            return False
        command, _ = self.parse_command(message)
        return command in SLASH_COMMANDS or command in AI_MODELS
    
    async def execute(self, message: str, progress_callback: Callable = None) -> Dict[str, Any]:
        """Execute a slash command and return results."""
        command, prompt = self.parse_command(message)
        
        if not command:
            return {"success": False, "error": "No command found"}
        
        # Check if it's a help command
        if command in ("help", "commands"):
            return self._show_help()
        
        # Check if it's a chain command
        if command == "chain":
            return await self._execute_chain(prompt, progress_callback)
        
        # Check if it's a registered slash command
        if command in SLASH_COMMANDS:
            cmd_info = SLASH_COMMANDS[command]
            cmd_type = cmd_info["type"]
            
            handler = self._get_command_handler(cmd_type)
            if handler:
                if asyncio.iscoroutinefunction(handler):
                    return await handler(command, prompt, cmd_info)
                else:
                    return handler(command, prompt, cmd_info)
        
        # Check if it's a direct model call
        if command in AI_MODELS:
            return await self._call_model(command, prompt)
        
        return {"success": False, "error": f"Unknown command: /{command}"}

    def _get_command_handler(self, cmd_type: str):
        """Get the handler method for a command type."""
        handlers = {
            "code": self._generate_code_file,
            "document": self._generate_document,
            "data": self._generate_data_file,
            "spreadsheet": self._generate_spreadsheet,
            "media": self._generate_media,
            "assistant": self._execute_assistant,
            "knowledge": self._execute_knowledge,
            "awareness": self._execute_awareness,
            "printify": self._execute_printify,
            "shopify": self._execute_shopify,
            "browser": self._execute_browser_command,
            "task_queue": self._execute_task_queue,
            "workflow": self._execute_workflow,
            "playground": self._execute_playground,
        }
        return handlers.get(cmd_type)
    
    def _show_help(self) -> Dict[str, Any]:
        """Generate help message showing all available commands organized by category."""
        # Use the new categorized structure
        help_text = """## ðŸš€ Otto Slash Commands

Type a command followed by your prompt, e.g., `/image a sunset over mountains`

### ðŸŽ¨ Media Generation
| Command | Description |
|---------|-------------|
| `/image` | Generate AI images |
| `/video` | Generate video clips |
| `/music` | Generate music tracks |
| `/speak` `/tts` | Text-to-speech |
| `/sound` `/sfx` | Sound effects |
| `/3d` | Generate 3D models |

### ðŸ“„ Documents
`/pdf` `/doc` `/docx` `/txt` `/rtf` `/md`

### ðŸ“Š Data Files
`/json` `/csv` `/xml` `/yaml` `/yml` `/toml` `/ini` `/env` `/xlsx` `/xls` `/sql`

### ðŸ’» Programming Languages
`/python` `/js` `/ts` `/java` `/c` `/cpp` `/go` `/rs` `/rb` `/php` `/swift` `/kt` `/scala` `/dart` `/lua` `/r` `/pl`

### ðŸŒ Web Development
`/html` `/css` `/scss` `/less` `/jsx` `/tsx` `/vue` `/svelte` `/astro` `/graphql`

### ðŸ“œ Scripts & Config
`/sh` `/bash` `/zsh` `/ps1` `/bat` `/dockerfile` `/makefile` `/nginx` `/gitignore`

### ðŸ¤– Direct AI Models
**Image:** `/flux` `/flux-dev` `/sdxl` `/imagen4` `/ideogram` `/bria`
**Video:** `/kling` `/minimax-video` `/luma`
**Music:** `/musicgen` `/lyria` `/stable-audio`
**3D:** `/hunyuan3d` `/rodin` `/luciddreamer`

### ðŸ§  AI Assistants (NEW!)
| Command | Description |
|---------|-------------|
| `/brand` | Generate brand colors, voice, names, full kit |
| `/outreach` | Generate emails, DMs, score leads |
| `/campaign` | Create campaign with goals, A/B variants |
| `/design` | Enhanced product design prompts |
| `/content` | Generate viral hooks, scored content |

### ðŸ“š Knowledge Base (NEW!)
| Command | Description |
|---------|-------------|
| `/learn` `/remember` | Add information to memory |
| `/recall` | Search knowledge base |
| `/knowledge` | Show knowledge base status |
| `/status` | Show current app state |
| `/forget` | Clear knowledge base |

### ðŸ“‹ Task Queue
| Command | Description |
|---------|-------------|
| `/task` | Create a multi-step task plan |
| `/queue` | View task queue status |
| `/batch` | Create batch operations (items -> action) |
| `/schedule` | Schedule recurring tasks |

### ðŸ”§ Workflows
| Command | Description |
|---------|-------------|
| `/workflow` | Create an automated workflow |
| `/run` | Execute a saved workflow |
| `/pipeline` | Create step -> step -> step pipeline |
| `/save` | Save workflow to file |
| `/load` | Load workflow from file |

### ðŸŽ® Playground
| Command | Description |
|---------|-------------|
| `/play` | Quick experiments (image/video/text/compare) |
| `/chain` | Chain commands: `/chain /image cat \| /video animate` |
| `/experiment` | List models or benchmark them |

### ðŸ”— Other
- `/help` - Show this help

---
**Examples:**
```
/image futuristic city at sunset
/python flask REST API with authentication
/brand modern tech startup for AI tools
/task create full marketing campaign for product launch
/workflow email newsletter automation
/pipeline logo design -> animate -> post to twitter
/play compare sunset landscape
/chain /image cat | /video animate
/batch img1.png, img2.png -> resize to 500x500
/schedule monday 9am -> send weekly analytics report
```
"""
        return {
            "success": True,
            "type": "help",
            "message": help_text,
            "artifacts": []
        }
    
    async def _execute_assistant(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute AI assistant commands (brand, outreach, campaign, design, content)."""
        assistant_type = cmd_info.get("assistant", command)
        
        try:
            if assistant_type == "brand_builder":
                return await self._ai_brand_builder(prompt)
            elif assistant_type == "outreach":
                return await self._ai_outreach_assistant(prompt)
            elif assistant_type == "campaign":
                return await self._ai_campaign_creator(prompt)
            elif assistant_type == "design":
                return await self._ai_design_studio(prompt)
            elif assistant_type == "content":
                return await self._ai_content_generator(prompt)
            else:
                return {"success": False, "error": f"Unknown assistant: {assistant_type}"}
        except Exception as e:
            logger.error(f"Assistant execution failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _ai_brand_builder(self, prompt: str) -> Dict[str, Any]:
        """AI Brand Builder - generates brand colors, voice, names, full kit."""
        brand_prompt = f"""You are an expert brand strategist. Based on this description, create a comprehensive brand kit:

BUSINESS/BRAND DESCRIPTION: {prompt}

Generate a complete brand kit in JSON format with:
{{
    "brand_names": ["3-5 creative brand name suggestions"],
    "taglines": ["3 catchy tagline options"],
    "color_palette": {{
        "primary": "#hexcolor - description",
        "secondary": "#hexcolor - description",
        "accent": "#hexcolor - description",
        "background": "#hexcolor",
        "text": "#hexcolor"
    }},
    "voice": {{
        "tone": "description of brand voice tone",
        "personality_traits": ["trait1", "trait2", "trait3"],
        "communication_style": "formal/casual/friendly/professional etc",
        "key_phrases": ["phrases the brand would use"],
        "words_to_avoid": ["words that don't fit the brand"]
    }},
    "target_audience": {{
        "demographics": "age, gender, location",
        "psychographics": "interests, values, lifestyle",
        "pain_points": ["what problems they have"],
        "goals": ["what they want to achieve"]
    }},
    "visual_style": {{
        "aesthetic": "modern/vintage/minimalist/bold etc",
        "imagery": "types of images that fit",
        "fonts_suggested": ["font style recommendations"]
    }}
}}

Return ONLY the JSON, no additional text."""

        try:
            result = self.replicate.generate_text(
                prompt=brand_prompt,
                max_tokens=1500,
                temperature=0.7
            )
            
            # Try to parse JSON
            json_match = re.search(r'\{[\s\S]*\}', result)
            if json_match:
                brand_kit = json.loads(json_match.group())
                
                # Store in session state for app use
                if brand_kit.get("brand_names"):
                    st.session_state.suggested_brand_names = brand_kit["brand_names"]
                if brand_kit.get("color_palette"):
                    st.session_state.suggested_brand_colors = brand_kit["color_palette"]
                if brand_kit.get("voice"):
                    st.session_state.suggested_brand_voice = brand_kit["voice"]
                
                # Format nice response
                response = "## ðŸŽ¨ AI Brand Kit Generated!\n\n"
                
                if brand_kit.get("brand_names"):
                    response += "### ðŸ’¡ Brand Name Ideas\n"
                    for name in brand_kit["brand_names"]:
                        response += f"- **{name}**\n"
                
                if brand_kit.get("taglines"):
                    response += "\n### âœ¨ Tagline Options\n"
                    for tagline in brand_kit["taglines"]:
                        response += f"- *\"{tagline}\"*\n"
                
                if brand_kit.get("color_palette"):
                    response += "\n### ðŸŽ¨ Color Palette\n"
                    colors = brand_kit["color_palette"]
                    for key, value in colors.items():
                        response += f"- **{key.title()}:** {value}\n"
                
                if brand_kit.get("voice"):
                    voice = brand_kit["voice"]
                    response += f"\n### ðŸ—£ï¸ Brand Voice\n"
                    response += f"**Tone:** {voice.get('tone', 'N/A')}\n"
                    response += f"**Style:** {voice.get('communication_style', 'N/A')}\n"
                    if voice.get("personality_traits"):
                        response += f"**Traits:** {', '.join(voice['personality_traits'])}\n"
                
                response += "\n*Brand kit saved to app settings. Use in Product Studio and Content Generator!*"
                
                return {
                    "success": True,
                    "type": "assistant",
                    "assistant": "brand_builder",
                    "message": response,
                    "data": brand_kit,
                    "artifacts": [{"type": "brand_kit", "data": brand_kit}]
                }
            else:
                return {
                    "success": True,
                    "type": "assistant",
                    "message": f"## ðŸŽ¨ Brand Suggestions\n\n{result}",
                    "artifacts": []
                }
                
        except Exception as e:
            logger.error(f"Brand builder failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _ai_outreach_assistant(self, prompt: str) -> Dict[str, Any]:
        """AI Outreach Assistant - generates emails, DMs, lead scoring."""
        outreach_prompt = f"""You are an expert cold outreach specialist. Generate outreach content for:

CONTEXT: {prompt}

Provide a comprehensive outreach package in this format:

## ðŸ“§ Cold Email Template
[A professional, personalized cold email with subject line, body, and CTA]

## ðŸ’¬ LinkedIn DM Script
[A concise, friendly LinkedIn message]

## ðŸ“± Twitter/X DM
[A short, engaging Twitter DM]

## ðŸŽ¯ Follow-up Sequence
1. [Day 3 follow-up]
2. [Day 7 follow-up]
3. [Day 14 final follow-up]

## ðŸ“Š Lead Scoring Criteria
- [What makes this a hot lead]
- [What makes this a warm lead]
- [What makes this a cold lead]

## ðŸ’¡ Personalization Tips
- [Specific ways to customize for each prospect]

Make the content professional but personal, with clear value propositions."""

        try:
            result = self.replicate.generate_text(
                prompt=outreach_prompt,
                max_tokens=1500,
                temperature=0.7
            )
            
            return {
                "success": True,
                "type": "assistant",
                "assistant": "outreach",
                "message": f"## ðŸŽ¯ AI Outreach Package\n\n{result}",
                "artifacts": [{"type": "outreach_content", "content": result}]
            }
            
        except Exception as e:
            logger.error(f"Outreach assistant failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _ai_campaign_creator(self, prompt: str) -> Dict[str, Any]:
        """AI Campaign Creator - creates campaign with goals, A/B variants, competitor analysis."""
        campaign_prompt = f"""You are a marketing campaign strategist. Create a comprehensive campaign for:

CAMPAIGN BRIEF: {prompt}

Generate a complete campaign plan:

## ðŸŽ¯ Campaign Goals
- Primary Goal: [specific, measurable goal]
- Secondary Goals: [2-3 supporting goals]
- KPIs to Track: [metrics]

## ðŸ“£ Core Message
**Headline:** [attention-grabbing headline]
**Value Proposition:** [clear benefit statement]
**Call to Action:** [specific CTA]

## ðŸ”„ A/B Testing Variants

### Variant A (Control)
- Headline: [version A]
- CTA: [version A]
- Visual Style: [description]

### Variant B (Test)
- Headline: [version B]
- CTA: [version B]
- Visual Style: [description]

## ðŸ“… Content Calendar (7 days)
| Day | Platform | Content Type | Caption/Copy |
|-----|----------|--------------|--------------|
[7 days of planned content]

## ðŸŽ¨ Visual Assets Needed
1. [asset 1 description]
2. [asset 2 description]
3. [asset 3 description]

## ðŸ’° Budget Allocation
- [channel 1]: X%
- [channel 2]: X%
- [channel 3]: X%

## ðŸ“Š Success Metrics
- [metric 1]: target value
- [metric 2]: target value"""

        try:
            result = self.replicate.generate_text(
                prompt=campaign_prompt,
                max_tokens=2000,
                temperature=0.6
            )
            
            # Store campaign in session
            st.session_state.generated_campaign = result
            
            return {
                "success": True,
                "type": "assistant",
                "assistant": "campaign",
                "message": f"## ðŸš€ AI Campaign Plan\n\n{result}\n\n*Campaign saved! Use Campaign Generator page to execute.*",
                "artifacts": [{"type": "campaign_plan", "content": result}]
            }
            
        except Exception as e:
            logger.error(f"Campaign creator failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _ai_design_studio(self, prompt: str) -> Dict[str, Any]:
        """AI Design Studio - enhanced product design prompts."""
        design_prompt = f"""You are a professional product designer and prompt engineer. Enhance this design request:

ORIGINAL REQUEST: {prompt}

Provide:

## ðŸŽ¨ Enhanced Design Prompt
[A detailed, professional prompt optimized for AI image generation. Include style, colors, composition, lighting, mood]

## ðŸ”¥ Trending Design Ideas
1. [Trending idea 1 based on current design trends]
2. [Trending idea 2]
3. [Trending idea 3]

## ðŸ“ Design Variations
### Minimalist Version
[Prompt for minimal design]

### Bold/Vibrant Version
[Prompt for bold design]

### Vintage/Retro Version
[Prompt for vintage design]

## ðŸŽ¯ Product Mockup Suggestions
- Best for: [t-shirt, mug, poster, etc.]
- Color scheme: [recommended colors]
- Placement: [where on product]

## ðŸ’¡ Style Keywords to Include
[List of professional keywords to enhance any prompt]"""

        try:
            result = self.replicate.generate_text(
                prompt=design_prompt,
                max_tokens=1200,
                temperature=0.7
            )
            
            return {
                "success": True,
                "type": "assistant",
                "assistant": "design",
                "message": f"## ðŸŽ¨ AI Design Studio\n\n{result}",
                "artifacts": [{"type": "design_guide", "content": result}]
            }
            
        except Exception as e:
            logger.error(f"Design studio failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _ai_content_generator(self, prompt: str) -> Dict[str, Any]:
        """AI Content Generator - generates viral hooks, scored content, variations."""
        content_prompt = f"""You are a viral content strategist. Create high-performing content for:

TOPIC/PRODUCT: {prompt}

Generate content optimized for engagement:

## ðŸŽ£ Viral Hooks (Opening Lines)
1. [Hook using curiosity gap]
2. [Hook using bold statement]
3. [Hook using question]
4. [Hook using story]
5. [Hook using controversy]

## ðŸ“± Platform-Specific Content

### Twitter/X (280 chars)
[Tweet with hashtags]
**Viral Score: X/10** | **Best Time:** [time]

### Instagram Caption
[Caption with emojis and hashtags]
**Engagement Score: X/10** | **Best Time:** [time]

### LinkedIn Post
[Professional post]
**Reach Score: X/10** | **Best Time:** [time]

### TikTok Script
[15-30 second script]
**Trend Score: X/10** | **Best Sound:** [suggestion]

## ðŸ“Š Content Scoring
| Platform | Engagement Potential | Virality Score | Timing |
|----------|---------------------|----------------|--------|
[Scores for each platform]

## ðŸ”„ A/B Variations
**Version A (Emotional):** [content variation]
**Version B (Logical):** [content variation]

## #ï¸âƒ£ Hashtag Strategy
Trending: [5 trending hashtags]
Niche: [5 niche hashtags]
Branded: [2-3 branded hashtag ideas]"""

        try:
            result = self.replicate.generate_text(
                prompt=content_prompt,
                max_tokens=1500,
                temperature=0.8
            )
            
            return {
                "success": True,
                "type": "assistant",
                "assistant": "content",
                "message": f"## âœ¨ AI Content Package\n\n{result}",
                "artifacts": [{"type": "content_package", "content": result}]
            }
            
        except Exception as e:
            logger.error(f"Content generator failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _execute_knowledge(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute knowledge base commands (learn, remember, recall, forget, knowledge)."""
        action = cmd_info.get("action", "status")
        kb = get_knowledge_base()
        
        try:
            if action == "add":
                if not prompt:
                    return {"success": False, "error": "Please provide information to remember. Usage: /learn [information]"}
                result = kb.add_fact(prompt)
                return {
                    "success": True,
                    "type": "knowledge",
                    "message": f"âœ… Added to knowledge base:\n\n*\"{prompt[:100]}...\"*\n\nI'll remember this for future conversations.",
                    "artifacts": []
                }
            
            elif action == "recall":
                if not prompt:
                    # Show all knowledge
                    summary = kb.get_context_summary()
                    return {
                        "success": True,
                        "type": "knowledge",
                        "message": f"## ðŸ“š Knowledge Base Summary\n\n{summary}",
                        "artifacts": []
                    }
                else:
                    result = kb.recall(prompt, self.replicate)
                    response = f"## ðŸ” Search Results for: *{prompt}*\n\n"
                    
                    if result["total_matches"] == 0:
                        response += "No matches found in knowledge base."
                    else:
                        if result["results"]["facts"]:
                            response += "### ðŸ“ Facts\n"
                            for fact in result["results"]["facts"]:
                                response += f"- {fact['fact']}\n"
                        
                        if result["results"]["documents"]:
                            response += "\n### ðŸ“„ Documents\n"
                            for doc in result["results"]["documents"]:
                                response += f"- **{doc['filename']}**: {doc['summary'][:100]}...\n"
                        
                        if result["results"]["images"]:
                            response += "\n### ðŸ–¼ï¸ Images\n"
                            for img in result["results"]["images"]:
                                response += f"- **{img['filename']}**: {img['analysis'][:100]}...\n"
                    
                    return {
                        "success": True,
                        "type": "knowledge",
                        "message": response,
                        "artifacts": []
                    }
            
            elif action == "clear":
                category = prompt.lower() if prompt else "all"
                result = kb.clear_memory(category)
                return {
                    "success": True,
                    "type": "knowledge",
                    "message": f"ðŸ—‘ï¸ Cleared {category} from knowledge base.",
                    "artifacts": []
                }
            
            elif action == "status":
                summary = kb.get_context_summary()
                stats = f"""## ðŸ“Š Knowledge Base Status

**Facts Stored:** {len(kb.facts)}
**Documents:** {len(kb.documents)}
**Images:** {len(kb.images)}

### Current Knowledge:
{summary}
"""
                return {
                    "success": True,
                    "type": "knowledge",
                    "message": stats,
                    "artifacts": []
                }
            
            return {"success": False, "error": f"Unknown knowledge action: {action}"}
            
        except Exception as e:
            logger.error(f"Knowledge command failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _execute_awareness(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute app awareness commands (status)."""
        try:
            state_summary = OttoAppAwareness.get_state_summary()
            full_state = OttoAppAwareness.get_current_state()
            
            response = f"""## ðŸ“ Current App State

{state_summary}

### ðŸ“Š Detailed State
```json
{json.dumps(full_state, indent=2, default=str)}
```

*I'm aware of everything happening in the app and can help you with any task!*
"""
            
            return {
                "success": True,
                "type": "awareness",
                "message": response,
                "state": full_state,
                "artifacts": []
            }
            
        except Exception as e:
            logger.error(f"Awareness command failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _execute_printify(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute Printify integration commands."""
        action = cmd_info.get("action", "menu")
        
        # Get Printify API from session state or environment
        printify_key = st.session_state.get("printify_api_key") or os.getenv("PRINTIFY_API_KEY")
        if not printify_key and action != "menu":
            return {
                "success": False,
                "error": "âŒ Printify API key not configured. Add it in Settings â†’ Integrations.",
                "artifacts": []
            }
        
        try:
            if action == "menu":
                menu = """## ðŸ›ï¸ Printify Commands

| Command | Description | Example |
|---------|-------------|---------|
| `/upload [image_url]` | Upload design to Printify | `/upload https://...` |
| `/product [type] [title]` | Create product from last upload | `/product mug Coffee Lovers Design` |
| `/publish [product_id]` | Publish product to store | `/publish 12345` |
| `/mockup [product_id]` | Get product mockup images | `/mockup 12345` |
| `/blueprint [search]` | Search product blueprints | `/blueprint hoodie` |
| `/shop` | Show shop info and products | `/shop` |

### ðŸš€ Quick Workflow
```
/image coffee cup design â†’ /upload â†’ /product mug Coffee Design â†’ /publish
```
"""
                return {"success": True, "type": "printify", "message": menu, "artifacts": []}
            
            # Import Printify API
            from printify import PrintifyAPI
            api = PrintifyAPI(printify_key)
            
            if action == "upload":
                # Upload image to Printify
                if not prompt:
                    # Check for recent generated image
                    if st.session_state.get("last_generated_image"):
                        image_url = st.session_state.last_generated_image
                    else:
                        return {"success": False, "error": "Please provide an image URL or generate an image first with /image"}
                else:
                    image_url = prompt.strip()
                
                st.info("ðŸ“¤ Uploading to Printify...")
                import requests
                response = requests.get(image_url, timeout=60)
                response.raise_for_status()
                image_bytes = response.content
                
                filename = f"otto_design_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                upload_id = api.upload_image(image_bytes, filename)
                
                # Store for later use
                st.session_state.last_printify_upload_id = upload_id
                
                return {
                    "success": True,
                    "type": "printify",
                    "message": f"âœ… **Image Uploaded to Printify!**\n\n**Upload ID:** `{upload_id}`\n\nUse `/product [type] [title]` to create a product.",
                    "data": {"upload_id": upload_id},
                    "artifacts": [{"type": "printify_upload", "id": upload_id}]
                }
            
            elif action == "create":
                # Create product
                parts = prompt.split(' ', 1) if prompt else []
                product_type = parts[0] if parts else "mug"
                title = parts[1] if len(parts) > 1 else f"Otto Design - {product_type.title()}"
                
                upload_id = st.session_state.get("last_printify_upload_id")
                if not upload_id:
                    return {"success": False, "error": "No design uploaded. Use /upload first or /image to generate."}
                
                st.info(f"ðŸ›ï¸ Creating {product_type} product...")
                
                # Find blueprint
                blueprint_id = api.find_blueprint(product_type)
                provider_id, variant_id, variant_info = api.get_provider_and_variant(blueprint_id)
                
                # Get shop
                shops = api.get_shops()
                if not shops:
                    return {"success": False, "error": "No Printify shops found"}
                shop_id = str(shops[0]["id"])
                
                # Create product
                product_data = {
                    "title": title,
                    "description": f"<p>AI-generated design by Otto Mate</p><p>{title}</p>",
                    "blueprint_id": blueprint_id,
                    "print_provider_id": provider_id,
                    "variants": [{"id": variant_id, "price": 2500, "is_enabled": True}],
                    "print_areas": [{
                        "variant_ids": [variant_id],
                        "placeholders": [{
                            "position": "front",
                            "images": [{"id": upload_id, "x": 0.5, "y": 0.5, "scale": 0.85, "angle": 0}]
                        }]
                    }]
                }
                
                product = api.create_product(shop_id, product_data)
                product_id = product.get("id")
                st.session_state.last_printify_product_id = product_id
                
                return {
                    "success": True,
                    "type": "printify",
                    "message": f"âœ… **Product Created!**\n\n**Title:** {title}\n**Product ID:** `{product_id}`\n**Type:** {product_type}\n\nUse `/publish` to make it live!",
                    "data": {"product_id": product_id, "title": title},
                    "artifacts": [{"type": "printify_product", "id": product_id}]
                }
            
            elif action == "publish":
                product_id = prompt.strip() if prompt else st.session_state.get("last_printify_product_id")
                if not product_id:
                    return {"success": False, "error": "No product ID. Create a product first with /product"}
                
                shops = api.get_shops()
                shop_id = str(shops[0]["id"])
                
                st.info("ðŸš€ Publishing product...")
                result = api.publish_product(shop_id, str(product_id))
                
                return {
                    "success": True,
                    "type": "printify",
                    "message": f"âœ… **Product Published!**\n\nProduct ID `{product_id}` is now live on your store!",
                    "data": {"product_id": product_id, "published": True},
                    "artifacts": []
                }
            
            elif action == "mockup":
                product_id = prompt.strip() if prompt else st.session_state.get("last_printify_product_id")
                if not product_id:
                    return {"success": False, "error": "No product ID provided"}
                
                shops = api.get_shops()
                shop_id = str(shops[0]["id"])
                
                mockups = api.get_all_product_mockups(shop_id, str(product_id))
                
                if mockups:
                    response = f"## ðŸ–¼ï¸ Product Mockups\n\n**Product ID:** `{product_id}`\n\n"
                    artifacts = []
                    for i, m in enumerate(mockups[:5]):  # Limit to 5
                        response += f"**Mockup {i+1}:** {'(Default)' if m['is_default'] else ''}\n"
                        artifacts.append({"type": "image", "url": m['url']})
                    
                    return {"success": True, "type": "printify", "message": response, "artifacts": artifacts}
                else:
                    return {"success": False, "error": "No mockups found for this product"}
            
            elif action == "blueprint":
                search = prompt.strip() if prompt else ""
                blueprints = api.get_blueprints()
                
                matches = [bp for bp in blueprints if search.lower() in bp.get("title", "").lower()][:10]
                
                response = f"## ðŸ“‹ Blueprint Search: '{search}'\n\n"
                if matches:
                    response += "| ID | Title | Description |\n|---|---|---|\n"
                    for bp in matches:
                        response += f"| {bp['id']} | {bp['title']} | {bp.get('description', '')[:50]} |\n"
                else:
                    response += "No blueprints found matching your search."
                
                return {"success": True, "type": "printify", "message": response, "artifacts": []}
            
            elif action == "shop":
                shops = api.get_shops()
                if shops:
                    shop = shops[0]
                    products = api.get_shop_products(str(shop["id"]), limit=10)
                    
                    response = f"""## ðŸª Printify Shop

**Shop Name:** {shop.get('title', 'N/A')}
**Shop ID:** {shop.get('id')}
**Sales Channel:** {shop.get('sales_channel_properties', {}).get('url', 'N/A')}

### Recent Products ({len(products)} shown)
"""
                    for p in products[:5]:
                        response += f"- **{p.get('title', 'Untitled')}** (ID: {p.get('id')})\n"
                    
                    return {"success": True, "type": "printify", "message": response, "artifacts": []}
                else:
                    return {"success": False, "error": "No shops found"}
            
            return {"success": False, "error": f"Unknown Printify action: {action}"}
            
        except Exception as e:
            logger.error(f"Printify command failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _execute_shopify(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute Shopify integration commands."""
        action = cmd_info.get("action", "menu")
        
        # Get Shopify credentials
        shop_url = st.session_state.get("shopify_store_url") or os.getenv("SHOPIFY_SHOP_URL")
        access_token = st.session_state.get("shopify_access_token") or os.getenv("SHOPIFY_ACCESS_TOKEN")
        
        if (not shop_url or not access_token) and action != "menu":
            return {
                "success": False,
                "error": "âŒ Shopify not configured. Add credentials in Settings â†’ Integrations.",
                "artifacts": []
            }
        
        try:
            if action == "menu":
                menu = """## ðŸª Shopify Commands

| Command | Description | Example |
|---------|-------------|---------|
| `/blog [topic]` | Create SEO-optimized blog post | `/blog benefits of coffee mugs` |
| `/store` | Show store info and stats | `/store` |
| `/seo [product]` | Generate SEO content | `/seo custom t-shirt` |
| `/inventory` | Check inventory levels | `/inventory` |
| `/analytics` | Get store analytics | `/analytics` |

### ðŸš€ Content Marketing Workflow
```
/brand tech startup â†’ /blog AI in business â†’ /content social posts for blog
```
"""
                return {"success": True, "type": "shopify", "message": menu, "artifacts": []}
            
            # Import Shopify API
            from shopify_service import ShopifyAPI
            api = ShopifyAPI(shop_url=shop_url, access_token=access_token)
            
            if action == "blog":
                if not prompt:
                    return {"success": False, "error": "Please provide a blog topic. Example: /blog benefits of morning coffee"}
                
                st.info("âœï¸ Generating SEO-optimized blog post...")
                
                # Generate blog content with AI
                blog_prompt = f"""Write an SEO-optimized blog post about: {prompt}

Requirements:
- Compelling title (include primary keyword)
- Meta description (155 characters)
- Introduction hook
- 3-5 subheadings with content
- Bullet points where appropriate
- Call-to-action
- 600-800 words
- HTML formatted

Output format:
TITLE: [title]
META: [meta description]
CONTENT:
[HTML content]"""

                blog_content = self.replicate.generate_text(
                    prompt=blog_prompt,
                    max_tokens=2000,
                    temperature=0.7
                )
                
                # Parse the generated content
                lines = blog_content.split('\n')
                title = prompt.title()
                meta = ""
                content = blog_content
                
                for i, line in enumerate(lines):
                    if line.startswith('TITLE:'):
                        title = line.replace('TITLE:', '').strip()
                    elif line.startswith('META:'):
                        meta = line.replace('META:', '').strip()
                    elif line.startswith('CONTENT:'):
                        content = '\n'.join(lines[i+1:])
                        break
                
                # Clean up content
                if not content.strip().startswith('<'):
                    content = f"<p>{content}</p>"
                
                # Publish to Shopify
                result = api.create_blog_post(
                    title=title,
                    body_html=content,
                    author="Otto Mate AI",
                    tags=["AI Generated", "SEO", prompt.split()[0] if prompt else "blog"],
                    published=True
                )
                
                if result:
                    return {
                        "success": True,
                        "type": "shopify",
                        "message": f"âœ… **Blog Post Published!**\n\n**Title:** {title}\n**URL:** {result.get('url', 'N/A')}\n**ID:** {result.get('id')}\n\n**Meta Description:** {meta[:100]}...",
                        "data": {"article_id": result.get("id"), "url": result.get("url")},
                        "artifacts": [{"type": "shopify_blog", "id": result.get("id"), "url": result.get("url")}]
                    }
                else:
                    return {"success": False, "error": "Failed to create blog post"}
            
            elif action == "store":
                if api.test_connection():
                    products = api.get_products(limit=10)
                    blogs = api.get_blogs()
                    
                    response = f"""## ðŸª Shopify Store

**Store URL:** {shop_url}
**Products:** {len(products)} shown
**Blogs:** {len(blogs)}

### Recent Products
"""
                    for p in products[:5]:
                        response += f"- **{p.get('title', 'Untitled')}** - ${p.get('variants', [{}])[0].get('price', 'N/A')}\n"
                    
                    return {"success": True, "type": "shopify", "message": response, "artifacts": []}
                else:
                    return {"success": False, "error": "Could not connect to Shopify store"}
            
            elif action == "seo":
                if not prompt:
                    return {"success": False, "error": "Please provide a product name. Example: /seo custom coffee mug"}
                
                seo_prompt = f"""Generate comprehensive SEO content for this product: {prompt}

Provide:
1. SEO Title (60 chars max)
2. Meta Description (155 chars max)
3. Product Description (200 words, HTML)
4. 10 SEO Keywords
5. 5 Long-tail Keywords
6. Alt Text for Images
7. URL Slug suggestion"""

                seo_content = self.replicate.generate_text(
                    prompt=seo_prompt,
                    max_tokens=1000,
                    temperature=0.6
                )
                
                return {
                    "success": True,
                    "type": "shopify",
                    "message": f"## ðŸ” SEO Content for: {prompt}\n\n{seo_content}",
                    "artifacts": [{"type": "text", "content": seo_content}]
                }
            
            elif action == "analytics":
                analytics = api.get_comprehensive_analytics() if hasattr(api, 'get_comprehensive_analytics') else {}
                
                response = f"""## ðŸ“Š Shopify Analytics

**Store:** {shop_url}

{json.dumps(analytics, indent=2) if analytics else 'Analytics data not available. Check API permissions.'}
"""
                return {"success": True, "type": "shopify", "message": response, "artifacts": []}
            
            elif action == "inventory":
                products = api.get_products(limit=50)
                
                response = "## ðŸ“¦ Inventory Status\n\n| Product | Variants | Status |\n|---|---|---|\n"
                for p in products[:20]:
                    variants = p.get('variants', [])
                    total_inventory = sum(v.get('inventory_quantity', 0) for v in variants)
                    status = "âœ… In Stock" if total_inventory > 0 else "âŒ Out of Stock"
                    response += f"| {p.get('title', 'Untitled')[:30]} | {len(variants)} | {status} |\n"
                
                return {"success": True, "type": "shopify", "message": response, "artifacts": []}
            
            return {"success": False, "error": f"Unknown Shopify action: {action}"}
            
        except Exception as e:
            logger.error(f"Shopify command failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _execute_browser_command(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute enhanced browser automation commands."""
        action = cmd_info.get("action", "browse")
        
        # Check for Anthropic key (required for browser automation)
        anthropic_key = os.getenv('ANTHROPIC_API_KEY')
        if not anthropic_key and action not in ["menu"]:
            return {
                "success": False,
                "error": "âŒ Browser automation requires ANTHROPIC_API_KEY. Add it to your .env file.",
                "artifacts": []
            }
        
        try:
            if action == "browse":
                if not prompt:
                    return {"success": False, "error": "Please provide a URL and task. Example: /browse https://example.com find pricing"}
                
                # Parse URL and task
                parts = prompt.split(' ', 1)
                url = parts[0] if parts else ""
                task = parts[1] if len(parts) > 1 else "Extract main content and summarize"
                
                st.info(f"ðŸŒ Browsing {url}...")
                
                from browser_use import Agent
                from langchain_anthropic import ChatAnthropic
                from pydantic import SecretStr
                
                llm = ChatAnthropic(
                    model_name="claude-sonnet-4-20250514",
                    api_key=SecretStr(anthropic_key),
                    temperature=0,
                    timeout=120
                )
                
                full_task = f"Go to {url} and {task}"
                agent = Agent(task=full_task, llm=llm)
                history = await agent.run()
                result_text = str(history.final_result()) if hasattr(history, 'final_result') else "Task completed"
                
                return {
                    "success": True,
                    "type": "browser",
                    "message": f"## ðŸŒ Browse Result\n\n**URL:** {url}\n**Task:** {task}\n\n**Result:**\n{result_text}",
                    "artifacts": [{"type": "text", "content": result_text}]
                }
            
            elif action == "scrape":
                if not prompt:
                    return {"success": False, "error": "Please provide URL and data to scrape. Example: /scrape https://example.com/products extract all product names and prices"}
                
                parts = prompt.split(' ', 1)
                url = parts[0]
                data_spec = parts[1] if len(parts) > 1 else "Extract all main content"
                
                st.info(f"ðŸ” Scraping {url}...")
                
                from browser_use import Agent
                from langchain_anthropic import ChatAnthropic
                from pydantic import SecretStr
                
                llm = ChatAnthropic(
                    model_name="claude-sonnet-4-20250514",
                    api_key=SecretStr(anthropic_key),
                    temperature=0,
                    timeout=120
                )
                
                scrape_task = f"Go to {url} and {data_spec}. Format the extracted data as structured JSON."
                agent = Agent(task=scrape_task, llm=llm)
                history = await agent.run()
                result_text = str(history.final_result()) if hasattr(history, 'final_result') else "{}"
                
                return {
                    "success": True,
                    "type": "browser",
                    "message": f"## ðŸ” Scraped Data\n\n**URL:** {url}\n\n```json\n{result_text}\n```",
                    "artifacts": [{"type": "text", "content": result_text}]
                }
            
            elif action == "post":
                if not prompt:
                    return {"success": False, "error": "Please specify platform and content. Example: /post twitter Just launched my new product!"}
                
                parts = prompt.split(' ', 1)
                platform = parts[0].lower()
                content = parts[1] if len(parts) > 1 else ""
                
                if not content:
                    return {"success": False, "error": "Please provide content to post"}
                
                st.info(f"ðŸ“± Posting to {platform}...")
                
                from browser_use import Agent
                from langchain_anthropic import ChatAnthropic
                from pydantic import SecretStr
                
                llm = ChatAnthropic(
                    model_name="claude-sonnet-4-20250514",
                    api_key=SecretStr(anthropic_key),
                    temperature=0,
                    timeout=180
                )
                
                if platform == "twitter":
                    task = f"Go to twitter.com, log in if needed, and post this tweet: {content}"
                elif platform == "linkedin":
                    task = f"Go to linkedin.com, log in if needed, and create a post with: {content}"
                elif platform == "facebook":
                    task = f"Go to facebook.com, log in if needed, and create a post with: {content}"
                else:
                    task = f"Go to {platform} and post: {content}"
                
                agent = Agent(task=task, llm=llm)
                history = await agent.run()
                result_text = str(history.final_result()) if hasattr(history, 'final_result') else "Posted"
                
                return {
                    "success": True,
                    "type": "browser",
                    "message": f"âœ… **Posted to {platform.title()}!**\n\n**Content:** {content[:100]}...\n\n**Result:** {result_text}",
                    "artifacts": []
                }
            
            elif action == "automate":
                if not prompt:
                    return {"success": False, "error": "Please describe the automation task. Example: /automate go to amazon.com and search for coffee mugs"}
                
                st.info(f"ðŸ¤– Running automation: {prompt[:50]}...")
                
                from browser_use import Agent
                from langchain_anthropic import ChatAnthropic
                from pydantic import SecretStr
                
                llm = ChatAnthropic(
                    model_name="claude-sonnet-4-20250514",
                    api_key=SecretStr(anthropic_key),
                    temperature=0,
                    timeout=180
                )
                
                agent = Agent(task=prompt, llm=llm)
                history = await agent.run()
                result_text = str(history.final_result()) if hasattr(history, 'final_result') else "Automation completed"
                
                return {
                    "success": True,
                    "type": "browser",
                    "message": f"## ðŸ¤– Automation Result\n\n**Task:** {prompt}\n\n**Result:**\n{result_text}",
                    "artifacts": [{"type": "text", "content": result_text}]
                }
            
            elif action == "screenshot":
                if not prompt:
                    return {"success": False, "error": "Please provide URL. Example: /screenshot https://example.com"}
                
                url = prompt.strip()
                st.info(f"ðŸ“¸ Taking screenshot of {url}...")
                
                from browser_use import Agent
                from langchain_anthropic import ChatAnthropic
                from pydantic import SecretStr
                
                llm = ChatAnthropic(
                    model_name="claude-sonnet-4-20250514",
                    api_key=SecretStr(anthropic_key),
                    temperature=0,
                    timeout=60
                )
                
                task = f"Go to {url} and take a screenshot. Describe what you see on the page."
                agent = Agent(task=task, llm=llm)
                history = await agent.run()
                result_text = str(history.final_result()) if hasattr(history, 'final_result') else "Screenshot taken"
                
                return {
                    "success": True,
                    "type": "browser",
                    "message": f"## ðŸ“¸ Screenshot Analysis\n\n**URL:** {url}\n\n**Description:**\n{result_text}",
                    "artifacts": []
                }
            
            return {"success": False, "error": f"Unknown browser action: {action}"}
            
        except ImportError as e:
            logger.error(f"Browser-Use import failed: {e}")
            return {"success": False, "error": "Browser-Use library not available. Install with: pip install browser-use langchain-anthropic"}
        except Exception as e:
            logger.error(f"Browser command failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _execute_task_queue(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute Task Queue commands - manage batch tasks, scheduling, and automation."""
        action = cmd_info.get("action", "create")
        
        try:
            # Import task queue engine
            from app.services.task_queue_engine import MultiStepPlanner, TaskQueueManager
            
            if action == "create":
                # /task create marketing campaign for new product
                if not prompt:
                    return {"success": False, "error": "Please describe the task. Example: /task create marketing materials for product launch"}
                
                # Initialize multi-step planner
                planner = MultiStepPlanner(api_client=self.replicate)
                
                # Create the plan from prompt
                plan = planner.create_plan(prompt)
                
                # Store in session state
                if 'otto_task_queue' not in st.session_state:
                    st.session_state.otto_task_queue = []
                
                task_entry = {
                    'id': datetime.now().strftime("%Y%m%d_%H%M%S"),
                    'prompt': prompt,
                    'plan': plan,
                    'status': 'queued',
                    'created': datetime.now().isoformat(),
                    'steps': [{'name': s['name'], 'status': 'pending'} for s in plan.get('steps', [])]
                }
                st.session_state.otto_task_queue.append(task_entry)
                
                # Format response
                steps_preview = "\n".join([f"  {i+1}. {s['name']}" for i, s in enumerate(plan.get('steps', [])[:5])])
                if len(plan.get('steps', [])) > 5:
                    steps_preview += f"\n  ... and {len(plan['steps']) - 5} more steps"
                
                return {
                    "success": True,
                    "type": "task_queue",
                    "message": f"""## ðŸ“‹ Task Created

**Task ID:** `{task_entry['id']}`
**Status:** Queued

**Plan Steps:**
{steps_preview}

Use `/queue status` to check progress or `/queue run` to execute.""",
                    "artifacts": [{"type": "task", "id": task_entry['id'], "plan": plan}]
                }
            
            elif action == "status":
                # /queue status - show all queued tasks
                queue = st.session_state.get('otto_task_queue', [])
                
                if not queue:
                    return {
                        "success": True,
                        "type": "task_queue",
                        "message": "ðŸ“‹ **Task Queue is empty**\n\nUse `/task <description>` to add a task.",
                        "artifacts": []
                    }
                
                # Format task list
                task_list = []
                for t in queue:
                    status_emoji = {"queued": "â³", "running": "ðŸ”„", "completed": "âœ…", "failed": "âŒ"}.get(t['status'], "âšª")
                    task_list.append(f"| `{t['id']}` | {t['prompt'][:40]}... | {status_emoji} {t['status']} |")
                
                tasks_table = "\n".join(task_list)
                
                return {
                    "success": True,
                    "type": "task_queue",
                    "message": f"""## ðŸ“‹ Task Queue Status

| ID | Description | Status |
|---|---|---|
{tasks_table}

**Commands:** `/queue run` â€¢ `/queue clear` â€¢ `/task <new task>`""",
                    "artifacts": []
                }
            
            elif action == "batch":
                # /batch image1, image2, image3 -> create variations
                if not prompt:
                    return {"success": False, "error": "Please provide batch items. Example: /batch logo1.png, logo2.png -> resize to 500x500"}
                
                # Parse batch format: items -> action
                parts = prompt.split('->')
                items = [i.strip() for i in parts[0].split(',')]
                action_desc = parts[1].strip() if len(parts) > 1 else "process"
                
                # Create batch tasks
                if 'otto_batch_tasks' not in st.session_state:
                    st.session_state.otto_batch_tasks = []
                
                batch_id = datetime.now().strftime("%Y%m%d_%H%M%S")
                batch = {
                    'id': batch_id,
                    'items': items,
                    'action': action_desc,
                    'status': 'ready',
                    'results': []
                }
                st.session_state.otto_batch_tasks.append(batch)
                
                return {
                    "success": True,
                    "type": "task_queue",
                    "message": f"""## ðŸ“¦ Batch Task Created

**Batch ID:** `{batch_id}`
**Items:** {len(items)} items queued
**Action:** {action_desc}

**Items:**
{chr(10).join([f'  â€¢ {item}' for item in items[:10]])}
{f'  ... and {len(items) - 10} more' if len(items) > 10 else ''}

Use `/batch run {batch_id}` to process all items.""",
                    "artifacts": []
                }
            
            elif action == "schedule":
                # /schedule 9am daily -> check analytics
                if not prompt:
                    return {"success": False, "error": "Please provide schedule and task. Example: /schedule 9am daily -> generate daily report"}
                
                # Parse: time pattern -> task
                parts = prompt.split('->')
                schedule_pattern = parts[0].strip()
                task_desc = parts[1].strip() if len(parts) > 1 else ""
                
                if not task_desc:
                    return {"success": False, "error": "Please specify what to schedule. Example: /schedule monday 9am -> post weekly newsletter"}
                
                # Store scheduled task
                if 'otto_scheduled_tasks' not in st.session_state:
                    st.session_state.otto_scheduled_tasks = []
                
                schedule_id = datetime.now().strftime("SCH_%Y%m%d_%H%M%S")
                scheduled = {
                    'id': schedule_id,
                    'pattern': schedule_pattern,
                    'task': task_desc,
                    'created': datetime.now().isoformat(),
                    'next_run': 'Calculating...',
                    'enabled': True
                }
                st.session_state.otto_scheduled_tasks.append(scheduled)
                
                return {
                    "success": True,
                    "type": "task_queue",
                    "message": f"""## â° Task Scheduled

**Schedule ID:** `{schedule_id}`
**When:** {schedule_pattern}
**Task:** {task_desc}

âœ… This task will run automatically according to the schedule.

**Manage:** `/schedule list` â€¢ `/schedule pause {schedule_id}` â€¢ `/schedule cancel {schedule_id}`""",
                    "artifacts": []
                }
            
            elif action == "run":
                # Run the next queued task or specified task
                queue = st.session_state.get('otto_task_queue', [])
                pending = [t for t in queue if t['status'] == 'queued']
                
                if not pending:
                    return {
                        "success": True,
                        "type": "task_queue",
                        "message": "ðŸ“‹ No pending tasks in queue. Use `/task <description>` to add one.",
                        "artifacts": []
                    }
                
                task = pending[0]
                task['status'] = 'running'
                
                st.info(f"ðŸ”„ Running task: {task['prompt'][:50]}...")
                
                # Execute the plan
                planner = MultiStepPlanner(api_client=self.replicate)
                results = await planner.execute_plan(task['plan'])
                
                task['status'] = 'completed'
                task['results'] = results
                
                return {
                    "success": True,
                    "type": "task_queue",
                    "message": f"""## âœ… Task Completed

**Task:** {task['prompt']}
**Duration:** {results.get('duration', 'N/A')}

{results.get('summary', 'Task executed successfully.')}""",
                    "artifacts": results.get('artifacts', [])
                }
            
            return {"success": False, "error": f"Unknown task queue action: {action}"}
            
        except Exception as e:
            logger.error(f"Task queue command failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _execute_workflow(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute Workflow commands - create, run, save, and manage workflows."""
        action = cmd_info.get("action", "create")
        
        try:
            if action == "create":
                # /workflow create product launch workflow
                if not prompt:
                    return {"success": False, "error": "Please describe the workflow. Example: /workflow new product launch automation"}
                
                # Use AI to generate workflow steps
                workflow_prompt = f"""Create a detailed workflow for: {prompt}

Return a JSON workflow with this structure:
{{
    "name": "Workflow Name",
    "description": "Brief description",
    "steps": [
        {{"id": 1, "name": "Step Name", "type": "action_type", "config": {{}}, "depends_on": []}},
    ],
    "triggers": ["manual"],
    "outputs": []
}}

Types can be: generate_image, generate_video, post_social, send_email, create_product, write_content, analyze_data
"""
                
                # Generate workflow
                response = self.replicate.generate_text(workflow_prompt, max_tokens=2000)
                
                # Parse JSON from response
                import json
                try:
                    # Extract JSON from response
                    json_match = response[response.find('{'):response.rfind('}')+1]
                    workflow = json.loads(json_match)
                except:
                    # Create basic workflow
                    workflow = {
                        "name": prompt[:50],
                        "description": prompt,
                        "steps": [{"id": 1, "name": "Execute main task", "type": "general", "config": {"prompt": prompt}}],
                        "triggers": ["manual"],
                        "outputs": []
                    }
                
                # Store workflow
                if 'otto_workflows' not in st.session_state:
                    st.session_state.otto_workflows = {}
                
                workflow_id = f"WF_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                workflow['id'] = workflow_id
                workflow['created'] = datetime.now().isoformat()
                st.session_state.otto_workflows[workflow_id] = workflow
                
                steps_list = "\n".join([f"  {s['id']}. **{s['name']}** ({s['type']})" for s in workflow['steps'][:8]])
                
                return {
                    "success": True,
                    "type": "workflow",
                    "message": f"""## ðŸ”§ Workflow Created

**ID:** `{workflow_id}`
**Name:** {workflow['name']}

**Steps:**
{steps_list}

**Commands:**
â€¢ `/run {workflow_id}` - Execute this workflow
â€¢ `/save {workflow_id} my_workflow` - Save to file
â€¢ `/workflow list` - View all workflows""",
                    "artifacts": [{"type": "workflow", "id": workflow_id, "workflow": workflow}]
                }
            
            elif action == "run":
                # /run WF_123 or /run "workflow name"
                if not prompt:
                    # List available workflows
                    workflows = st.session_state.get('otto_workflows', {})
                    if not workflows:
                        return {"success": False, "error": "No workflows available. Use `/workflow <description>` to create one."}
                    
                    wf_list = "\n".join([f"â€¢ `{wid}`: {w['name']}" for wid, w in list(workflows.items())[:10]])
                    return {
                        "success": True,
                        "type": "workflow",
                        "message": f"**Available Workflows:**\n{wf_list}\n\nUse `/run <workflow_id>` to execute.",
                        "artifacts": []
                    }
                
                # Find workflow
                workflows = st.session_state.get('otto_workflows', {})
                workflow = None
                
                if prompt.upper().startswith('WF_'):
                    workflow = workflows.get(prompt)
                else:
                    # Search by name
                    for wid, w in workflows.items():
                        if prompt.lower() in w['name'].lower():
                            workflow = w
                            break
                
                if not workflow:
                    return {"success": False, "error": f"Workflow '{prompt}' not found. Use `/workflow list` to see available workflows."}
                
                st.info(f"ðŸ”§ Running workflow: {workflow['name']}")
                
                # Execute each step
                results = []
                for step in workflow['steps']:
                    st.info(f"  Step {step['id']}: {step['name']}")
                    
                    # Execute based on type
                    step_result = {"step": step['name'], "status": "completed"}
                    
                    if step['type'] == 'generate_image':
                        result = await self._generate_media('image', step['config'].get('prompt', ''), {'model': 'flux'})
                        step_result['result'] = result
                    elif step['type'] == 'generate_video':
                        result = await self._generate_media('video', step['config'].get('prompt', ''), {'model': 'kling'})
                        step_result['result'] = result
                    elif step['type'] == 'write_content':
                        result = self.replicate.generate_text(step['config'].get('prompt', ''))
                        step_result['result'] = {"content": result}
                    else:
                        step_result['result'] = {"message": "Step executed"}
                    
                    results.append(step_result)
                
                return {
                    "success": True,
                    "type": "workflow",
                    "message": f"""## âœ… Workflow Completed

**{workflow['name']}**

**Results:**
{chr(10).join([f"  âœ“ {r['step']}" for r in results])}

All {len(results)} steps completed successfully.""",
                    "artifacts": [r.get('result', {}).get('artifacts', []) for r in results if r.get('result')]
                }
            
            elif action == "save":
                # /save WF_123 my_workflow
                if not prompt:
                    return {"success": False, "error": "Please specify workflow ID and filename. Example: /save WF_123 my_workflow"}
                
                parts = prompt.split()
                workflow_id = parts[0]
                filename = parts[1] if len(parts) > 1 else workflow_id
                
                workflows = st.session_state.get('otto_workflows', {})
                workflow = workflows.get(workflow_id)
                
                if not workflow:
                    return {"success": False, "error": f"Workflow '{workflow_id}' not found."}
                
                # Save to file
                import json
                filepath = self.output_dir / f"{filename}.workflow.json"
                with open(filepath, 'w') as f:
                    json.dump(workflow, f, indent=2, default=str)
                
                return {
                    "success": True,
                    "type": "workflow",
                    "message": f"âœ… Workflow saved to `{filepath}`",
                    "artifacts": [{"type": "file", "path": str(filepath)}]
                }
            
            elif action == "load":
                # /load my_workflow
                if not prompt:
                    return {"success": False, "error": "Please specify workflow file. Example: /load my_workflow"}
                
                import json
                
                # Try different paths
                filename = prompt if prompt.endswith('.json') else f"{prompt}.workflow.json"
                filepath = self.output_dir / filename
                
                if not filepath.exists():
                    filepath = Path(prompt)
                
                if not filepath.exists():
                    return {"success": False, "error": f"Workflow file not found: {prompt}"}
                
                with open(filepath) as f:
                    workflow = json.load(f)
                
                # Store in session
                if 'otto_workflows' not in st.session_state:
                    st.session_state.otto_workflows = {}
                
                workflow_id = workflow.get('id', f"WF_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
                st.session_state.otto_workflows[workflow_id] = workflow
                
                return {
                    "success": True,
                    "type": "workflow",
                    "message": f"âœ… Loaded workflow: **{workflow['name']}** (`{workflow_id}`)\n\nUse `/run {workflow_id}` to execute.",
                    "artifacts": []
                }
            
            elif action == "list":
                # /workflow list
                workflows = st.session_state.get('otto_workflows', {})
                
                if not workflows:
                    return {
                        "success": True,
                        "type": "workflow",
                        "message": "ðŸ“ No workflows yet. Use `/workflow <description>` to create one.",
                        "artifacts": []
                    }
                
                wf_rows = []
                for wid, w in workflows.items():
                    wf_rows.append(f"| `{wid}` | {w['name'][:30]} | {len(w['steps'])} steps |")
                
                return {
                    "success": True,
                    "type": "workflow",
                    "message": f"""## ðŸ”§ Your Workflows

| ID | Name | Steps |
|---|---|---|
{chr(10).join(wf_rows)}

**Commands:** `/run <id>` â€¢ `/workflow new <desc>` â€¢ `/save <id> <name>`""",
                    "artifacts": []
                }
            
            elif action == "pipeline":
                # /pipeline image -> video -> post
                if not prompt:
                    return {"success": False, "error": "Please describe the pipeline. Example: /pipeline generate logo -> animate -> post to twitter"}
                
                # Parse pipeline steps
                steps = [s.strip() for s in prompt.split('->')]
                
                # Create workflow from pipeline
                workflow_steps = []
                for i, step in enumerate(steps):
                    step_type = "general"
                    if any(w in step.lower() for w in ['image', 'logo', 'design', 'graphic']):
                        step_type = "generate_image"
                    elif any(w in step.lower() for w in ['video', 'animate', 'clip']):
                        step_type = "generate_video"
                    elif any(w in step.lower() for w in ['post', 'tweet', 'publish', 'share']):
                        step_type = "post_social"
                    elif any(w in step.lower() for w in ['email', 'send', 'newsletter']):
                        step_type = "send_email"
                    elif any(w in step.lower() for w in ['write', 'content', 'blog', 'article']):
                        step_type = "write_content"
                    
                    workflow_steps.append({
                        "id": i + 1,
                        "name": step,
                        "type": step_type,
                        "config": {"prompt": step},
                        "depends_on": [i] if i > 0 else []
                    })
                
                workflow = {
                    "name": f"Pipeline: {steps[0][:20]}...",
                    "description": prompt,
                    "steps": workflow_steps,
                    "triggers": ["manual"],
                    "outputs": []
                }
                
                # Store workflow
                if 'otto_workflows' not in st.session_state:
                    st.session_state.otto_workflows = {}
                
                workflow_id = f"PIPE_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                workflow['id'] = workflow_id
                st.session_state.otto_workflows[workflow_id] = workflow
                
                pipeline_visual = " â†’ ".join([f"**{s['name'][:15]}**" for s in workflow_steps])
                
                return {
                    "success": True,
                    "type": "workflow",
                    "message": f"""## ðŸ”„ Pipeline Created

**ID:** `{workflow_id}`

**Flow:**
{pipeline_visual}

**Steps:** {len(workflow_steps)} stages

Use `/run {workflow_id}` to execute the pipeline.""",
                    "artifacts": [{"type": "workflow", "id": workflow_id}]
                }
            
            return {"success": False, "error": f"Unknown workflow action: {action}"}
            
        except Exception as e:
            logger.error(f"Workflow command failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _execute_playground(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Execute Playground commands - experiment with models, chain outputs, and explore AI capabilities."""
        action = cmd_info.get("action", "experiment")
        
        try:
            if action == "experiment":
                # /play with flux create abstract art
                if not prompt:
                    return {
                        "success": True,
                        "type": "playground",
                        "message": """## ðŸŽ® Otto Playground

**Quick Experiments:**
â€¢ `/play image <prompt>` - Quick image generation
â€¢ `/play video <prompt>` - Quick video generation
â€¢ `/play text <prompt>` - Quick text generation
â€¢ `/play compare <prompt>` - Compare multiple models

**Advanced:**
â€¢ `/chain <step1> | <step2>` - Chain model outputs
â€¢ `/experiment models` - List all available models
â€¢ `/experiment benchmark <prompt>` - Benchmark models

**Examples:**
```
/play image cyberpunk cat
/play compare sunset landscape
/chain /image dog | /video animate
```""",
                        "artifacts": []
                    }
                
                # Parse play command
                parts = prompt.split(' ', 1)
                play_type = parts[0].lower()
                play_prompt = parts[1] if len(parts) > 1 else ""
                
                if play_type == 'image':
                    result = await self._generate_media('image', play_prompt, {'model': 'flux'})
                    return result
                elif play_type == 'video':
                    result = await self._generate_media('video', play_prompt, {'model': 'kling'})
                    return result
                elif play_type == 'text':
                    content = self.replicate.generate_text(play_prompt, max_tokens=1000)
                    return {
                        "success": True,
                        "type": "playground",
                        "message": f"## ðŸ“ Generated Text\n\n{content}",
                        "artifacts": [{"type": "text", "content": content}]
                    }
                elif play_type == 'compare':
                    # Generate with multiple models for comparison
                    st.info("ðŸ”„ Generating with multiple models for comparison...")
                    
                    models_to_compare = ['flux', 'sdxl']
                    results = []
                    
                    for model in models_to_compare:
                        try:
                            result = await self._generate_media('image', play_prompt, {'model': model})
                            if result.get('success'):
                                results.append({'model': model, 'result': result})
                        except:
                            pass
                    
                    return {
                        "success": True,
                        "type": "playground",
                        "message": f"## ðŸ”„ Model Comparison\n\n**Prompt:** {play_prompt}\n\n**Results from {len(results)} models:**",
                        "artifacts": [r['result'].get('artifacts', [{}])[0] for r in results if r['result'].get('artifacts')]
                    }
                else:
                    # Treat as general prompt
                    result = await self._generate_media('image', prompt, {'model': 'flux'})
                    return result
            
            elif action == "chain":
                # /chain /image cat | /video animate | /music add soundtrack
                if not prompt:
                    return {"success": False, "error": "Please provide a chain. Example: /chain /image cat | /video animate"}
                
                # Parse chain steps
                steps = [s.strip() for s in prompt.split('|')]
                
                st.info(f"â›“ï¸ Executing chain with {len(steps)} steps...")
                
                chain_results = []
                previous_output = None
                
                for i, step in enumerate(steps):
                    st.info(f"  Step {i+1}: {step[:30]}...")
                    
                    # Parse the step command
                    if step.startswith('/'):
                        cmd_parts = step[1:].split(' ', 1)
                        cmd = cmd_parts[0]
                        cmd_prompt = cmd_parts[1] if len(cmd_parts) > 1 else ""
                        
                        # If we have previous output, inject it
                        if previous_output and not cmd_prompt:
                            cmd_prompt = f"Based on previous: {previous_output}"
                        
                        # Execute the command
                        result = await self.execute(f"/{cmd} {cmd_prompt}")
                        chain_results.append({'step': step, 'result': result})
                        
                        # Save output for next step
                        if result.get('success'):
                            if result.get('artifacts'):
                                previous_output = result['artifacts'][0].get('url') or result['artifacts'][0].get('content')
                            elif result.get('message'):
                                previous_output = result['message'][:200]
                    else:
                        # Treat as text prompt
                        content = self.replicate.generate_text(step)
                        chain_results.append({'step': step, 'result': {'success': True, 'content': content}})
                        previous_output = content
                
                # Collect all artifacts
                all_artifacts = []
                for r in chain_results:
                    if r['result'].get('artifacts'):
                        all_artifacts.extend(r['result']['artifacts'])
                
                steps_summary = "\n".join([f"  {i+1}. {r['step'][:40]}... {'âœ…' if r['result'].get('success') else 'âŒ'}" for i, r in enumerate(chain_results)])
                
                return {
                    "success": True,
                    "type": "playground",
                    "message": f"""## â›“ï¸ Chain Completed

**Steps:**
{steps_summary}

**Generated {len(all_artifacts)} artifacts**""",
                    "artifacts": all_artifacts
                }
            
            elif action == "models":
                # /experiment models - list all available models
                models_by_category = {
                    "ðŸŽ¨ Image": ["flux", "flux-dev", "sdxl", "imagen4", "ideogram", "bria"],
                    "ðŸŽ¬ Video": ["kling", "minimax-video", "luma", "wan"],
                    "ðŸŽµ Music": ["musicgen", "lyria", "stable-audio"],
                    "ðŸ—£ï¸ Voice": ["xtts", "eleven", "cartesia"],
                    "ðŸ§Š 3D": ["hunyuan3d", "rodin", "luciddreamer"],
                    "ðŸ’¬ Text": ["llama3", "mistral", "claude"]
                }
                
                model_text = ""
                for category, models in models_by_category.items():
                    model_text += f"\n**{category}:** {', '.join([f'`{m}`' for m in models])}"
                
                return {
                    "success": True,
                    "type": "playground",
                    "message": f"""## ðŸ¤– Available AI Models
{model_text}

**Usage:**
â€¢ Direct model call: `/flux a sunset`
â€¢ With play: `/play image a sunset`
â€¢ In chains: `/chain /flux cat | /kling animate`""",
                    "artifacts": []
                }
            
            elif action == "benchmark":
                # /experiment benchmark <prompt>
                if not prompt:
                    return {"success": False, "error": "Please provide a prompt to benchmark. Example: /experiment benchmark a red rose"}
                
                st.info("ðŸƒ Running benchmark across models...")
                
                models = ['flux', 'sdxl']
                benchmark_results = []
                
                for model in models:
                    start_time = datetime.now()
                    try:
                        result = await self._generate_media('image', prompt, {'model': model})
                        duration = (datetime.now() - start_time).total_seconds()
                        benchmark_results.append({
                            'model': model,
                            'success': result.get('success', False),
                            'duration': f"{duration:.1f}s",
                            'result': result
                        })
                    except Exception as e:
                        benchmark_results.append({
                            'model': model,
                            'success': False,
                            'duration': 'Failed',
                            'error': str(e)
                        })
                
                results_table = "\n".join([f"| {r['model']} | {'âœ…' if r['success'] else 'âŒ'} | {r['duration']} |" for r in benchmark_results])
                
                return {
                    "success": True,
                    "type": "playground",
                    "message": f"""## ðŸƒ Benchmark Results

**Prompt:** {prompt}

| Model | Status | Duration |
|---|---|---|
{results_table}

**Winner:** {min([r for r in benchmark_results if r['success']], key=lambda x: float(x['duration'].replace('s','')), default={'model': 'N/A'})['model']}""",
                    "artifacts": [r['result'].get('artifacts', [{}])[0] for r in benchmark_results if r.get('result', {}).get('artifacts')]
                }
            
            return {"success": False, "error": f"Unknown playground action: {action}"}
            
        except Exception as e:
            logger.error(f"Playground command failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _generate_code_file(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Generate code/text files using AI."""
        ext = cmd_info["ext"]
        
        # Language-specific prompts
        lang_prompts = {
            # Python & Scripting
            "py": "Write complete, well-documented Python code",
            "rb": "Write clean Ruby code with proper conventions",
            "pl": "Write Perl code",
            "lua": "Write Lua code",
            "r": "Write R code for statistical computing",
            
            # Web
            "html": "Write complete, valid HTML5 with modern CSS styling",
            "css": "Write modern CSS with flexbox/grid where appropriate",
            "scss": "Write SCSS with proper nesting and variables",
            "less": "Write LESS stylesheet",
            "js": "Write clean, modern JavaScript (ES6+)",
            "ts": "Write TypeScript with proper types",
            "jsx": "Write React JSX component",
            "tsx": "Write React TypeScript component",
            "vue": "Write Vue.js single-file component",
            "svelte": "Write Svelte component",
            "astro": "Write Astro component",
            
            # Systems
            "c": "Write clean, well-commented C code",
            "cpp": "Write modern C++ code (C++17/20) with proper headers",
            "h": "Write C header file with proper guards",
            "hpp": "Write C++ header file",
            "go": "Write idiomatic Go code",
            "rs": "Write safe, idiomatic Rust code",
            "swift": "Write Swift code following Apple guidelines",
            "kt": "Write idiomatic Kotlin code",
            "java": "Write complete Java code with proper class structure",
            "scala": "Write Scala code",
            "dart": "Write Dart code",
            
            # Backend
            "php": "Write modern PHP code (PHP 8+)",
            "sql": "Write SQL code with proper formatting",
            "graphql": "Write GraphQL schema or query",
            "proto": "Write Protocol Buffers definition",
            
            # Scripts
            "bat": "Write a Windows batch script",
            "sh": "Write a Unix/Linux shell script (bash)",
            "bash": "Write a Bash script",
            "zsh": "Write a Zsh script",
            "ps1": "Write a PowerShell script",
            
            # Config
            "Dockerfile": "Write a Dockerfile with best practices",
            "Makefile": "Write a Makefile",
            "CMakeLists.txt": "Write CMake configuration",
            "conf": "Write Nginx configuration",
            ".htaccess": "Write Apache .htaccess file",
            ".gitignore": "Write .gitignore file",
            
            # Docs
            "md": "Write well-formatted Markdown",
            "txt": "Write clear, well-organized text",
            "log": "Generate realistic log entries",
        }
        
        lang_instruction = lang_prompts.get(ext, f"Generate {ext} content")
        
        generation_prompt = f"""{lang_instruction} for the following request:

REQUEST: {prompt}

Requirements:
1. The code/content must be complete and functional
2. Include helpful comments where appropriate
3. Follow best practices for the language/format
4. Output ONLY the raw code/content - no markdown code blocks or explanations

Begin output:"""

        try:
            content = self.replicate.generate_text(
                prompt=generation_prompt,
                max_tokens=4000,
                temperature=0.3
            )
            
            # Clean up the response
            content = content.strip()
            if content.startswith("```"):
                # Remove markdown code blocks
                lines = content.split('\n')
                content = '\n'.join(lines[1:-1] if lines[-1] == "```" else lines[1:])
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"otto_{command}_{timestamp}.{ext}"
            filepath = self.output_dir / filename
            
            # Write file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                "success": True,
                "type": "file",
                "file_type": ext,
                "filename": filename,
                "filepath": str(filepath),
                "content": content,
                "message": f"âœ… Generated {ext.upper()} file: {filename}",
                "artifacts": [{"type": "file", "path": str(filepath), "filename": filename, "ext": ext}]
            }
            
        except Exception as e:
            logger.error(f"Code generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _generate_document(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Generate document files (PDF, DOCX, DOC, RTF)."""
        ext = cmd_info["ext"]
        
        # First generate the content
        generation_prompt = f"""Write professional document content for:

REQUEST: {prompt}

Create well-structured content with:
- Clear title/heading
- Organized sections
- Professional language
- Appropriate formatting

Output the content in plain text with clear section breaks."""

        try:
            content = self.replicate.generate_text(
                prompt=generation_prompt,
                max_tokens=4000,
                temperature=0.4
            )
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"otto_doc_{timestamp}.{ext}"
            filepath = self.output_dir / filename
            
            if ext == "pdf":
                # Generate PDF using reportlab or fpdf
                try:
                    from fpdf import FPDF
                    
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    
                    # Split content into lines and add to PDF
                    for line in content.split('\n'):
                        pdf.multi_cell(0, 10, txt=line)
                    
                    pdf.output(str(filepath))
                    
                except ImportError:
                    # Fallback: Save as HTML that can be printed to PDF
                    html_content = f"""<!DOCTYPE html>
<html><head><title>Generated Document</title>
<style>body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; }}</style>
</head><body>
<pre style="white-space: pre-wrap;">{content}</pre>
</body></html>"""
                    filepath = self.output_dir / f"otto_doc_{timestamp}.html"
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                    return {
                        "success": True,
                        "type": "file",
                        "file_type": "html",
                        "filename": filepath.name,
                        "filepath": str(filepath),
                        "content": content,
                        "message": f"âœ… PDF library not available. Generated HTML file instead: {filepath.name}\n\nOpen and print to PDF.",
                        "artifacts": [{"type": "file", "path": str(filepath), "filename": filepath.name, "ext": "html"}]
                    }
            
            elif ext == "docx":
                try:
                    from docx import Document
                    
                    doc = Document()
                    
                    # Add content with basic formatting
                    for para in content.split('\n\n'):
                        doc.add_paragraph(para)
                    
                    doc.save(str(filepath))
                    
                except ImportError:
                    # Fallback to RTF
                    filepath = self.output_dir / f"otto_doc_{timestamp}.rtf"
                    rtf_content = r"{\rtf1\ansi " + content.replace('\n', r'\par ') + "}"
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(rtf_content)
                    return {
                        "success": True,
                        "type": "file",
                        "filename": filepath.name,
                        "filepath": str(filepath),
                        "content": content,
                        "message": f"âœ… python-docx not available. Generated RTF file instead: {filepath.name}",
                        "artifacts": [{"type": "file", "path": str(filepath), "filename": filepath.name, "ext": "rtf"}]
                    }
            
            elif ext in ["doc", "rtf"]:
                # Generate RTF format
                rtf_content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Arial;}}\f0\fs24 " + content.replace('\n', r'\par ') + "}"
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(rtf_content)
            
            return {
                "success": True,
                "type": "file",
                "file_type": ext,
                "filename": filename,
                "filepath": str(filepath),
                "content": content,
                "message": f"âœ… Generated {ext.upper()} document: {filename}",
                "artifacts": [{"type": "file", "path": str(filepath), "filename": filename, "ext": ext}]
            }
            
        except Exception as e:
            logger.error(f"Document generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _generate_data_file(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Generate data files (CSV, JSON, XML, YAML, INI)."""
        ext = cmd_info["ext"]
        
        format_instructions = {
            "csv": "Generate data in CSV format with a header row. Output ONLY the CSV data.",
            "json": "Generate valid JSON data. Output ONLY the raw JSON.",
            "xml": "Generate valid XML data with proper structure. Output ONLY the raw XML.",
            "yaml": "Generate valid YAML data. Output ONLY the raw YAML.",
            "ini": "Generate valid INI configuration file format. Output ONLY the raw INI content.",
        }
        
        generation_prompt = f"""{format_instructions.get(ext, "Generate the data")}

REQUEST: {prompt}

Output only the raw data in the correct format - no explanations or code blocks."""

        try:
            content = self.replicate.generate_text(
                prompt=generation_prompt,
                max_tokens=4000,
                temperature=0.2  # Lower temp for structured data
            )
            
            # Clean up
            content = content.strip()
            if content.startswith("```"):
                lines = content.split('\n')
                content = '\n'.join(lines[1:-1] if lines[-1].startswith("```") else lines[1:])
            
            # Validate and format
            if ext == "json":
                try:
                    # Validate JSON
                    parsed = json.loads(content)
                    content = json.dumps(parsed, indent=2)
                except json.JSONDecodeError:
                    # Try to extract JSON from response
                    json_match = re.search(r'[\[{][\s\S]*[\]}]', content)
                    if json_match:
                        content = json_match.group()
                        parsed = json.loads(content)
                        content = json.dumps(parsed, indent=2)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"otto_data_{timestamp}.{ext}"
            filepath = self.output_dir / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                "success": True,
                "type": "file",
                "file_type": ext,
                "filename": filename,
                "filepath": str(filepath),
                "content": content,
                "message": f"âœ… Generated {ext.upper()} file: {filename}",
                "artifacts": [{"type": "file", "path": str(filepath), "filename": filename, "ext": ext}]
            }
            
        except Exception as e:
            logger.error(f"Data file generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _generate_spreadsheet(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Generate spreadsheet files (XLS, XLSX)."""
        ext = cmd_info["ext"]
        
        # First generate CSV-like data
        generation_prompt = f"""Generate spreadsheet data for:

REQUEST: {prompt}

Output as CSV format with:
- First row as headers
- Properly formatted data in columns
- Meaningful sample data (at least 5-10 rows)

Output ONLY the CSV data - no explanations."""

        try:
            csv_content = self.replicate.generate_text(
                prompt=generation_prompt,
                max_tokens=4000,
                temperature=0.3
            )
            
            # Clean up
            csv_content = csv_content.strip()
            if csv_content.startswith("```"):
                lines = csv_content.split('\n')
                csv_content = '\n'.join(lines[1:-1] if lines[-1].startswith("```") else lines[1:])
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"otto_spreadsheet_{timestamp}.{ext}"
            filepath = self.output_dir / filename
            
            try:
                import pandas as pd
                import io
                
                # Parse CSV to DataFrame
                df = pd.read_csv(io.StringIO(csv_content))
                
                # Write to Excel
                if ext == "xlsx":
                    df.to_excel(str(filepath), index=False, engine='openpyxl')
                else:  # xls
                    # Fallback to CSV if xlwt not available
                    df.to_excel(str(filepath), index=False)
                    
            except ImportError:
                # Fallback to CSV
                filepath = self.output_dir / f"otto_spreadsheet_{timestamp}.csv"
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(csv_content)
                return {
                    "success": True,
                    "type": "file",
                    "filename": filepath.name,
                    "filepath": str(filepath),
                    "content": csv_content,
                    "message": f"âœ… Excel libraries not available. Generated CSV file instead: {filepath.name}",
                    "artifacts": [{"type": "file", "path": str(filepath), "filename": filepath.name, "ext": "csv"}]
                }
            
            return {
                "success": True,
                "type": "file",
                "file_type": ext,
                "filename": filename,
                "filepath": str(filepath),
                "content": csv_content,
                "message": f"âœ… Generated {ext.upper()} spreadsheet: {filename}",
                "artifacts": [{"type": "file", "path": str(filepath), "filename": filename, "ext": ext}]
            }
            
        except Exception as e:
            logger.error(f"Spreadsheet generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _generate_media(self, command: str, prompt: str, cmd_info: Dict) -> Dict[str, Any]:
        """Generate media files (image, video, music, speech, sound, 3d)."""
        model_key = cmd_info.get("model", command)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        try:
            if command == "image":
                # Use flux-fast for images
                st.info("ðŸŽ¨ Generating image...")
                image_url = self.replicate.generate_image(
                    prompt=prompt,
                    width=1024,
                    height=1024
                )
                
                # Save to campaigns folder
                filename = f"image_{timestamp}.png"
                saved_path = self._save_to_campaigns(filename, url=image_url)
                
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "image",
                    "url": image_url,
                    "filepath": saved_path,
                    "message": f"âœ… Generated image",
                    "artifacts": [{"type": "image", "url": image_url, "filepath": saved_path, "prompt": prompt}]
                }
            
            elif command == "video":
                # Generate image first, then video
                st.info("ðŸŽ¨ Generating reference image...")
                image_url = self.replicate.generate_image(prompt=prompt, width=1280, height=720)
                st.info("ðŸŽ¬ Generating video (this may take a minute)...")
                video_url = self.replicate.generate_video(
                    prompt=f"Gentle cinematic motion: {prompt}",
                    image_url=image_url
                )
                
                # Save to campaigns folder
                video_filename = f"video_{timestamp}.mp4"
                image_filename = f"video_ref_{timestamp}.png"
                video_path = self._save_to_campaigns(video_filename, url=video_url)
                image_path = self._save_to_campaigns(image_filename, url=image_url)
                
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "video",
                    "url": video_url,
                    "source_image": image_url,
                    "filepath": video_path,
                    "message": f"âœ… Generated video",
                    "artifacts": [
                        {"type": "video", "url": video_url, "filepath": video_path},
                        {"type": "image", "url": image_url, "filepath": image_path}
                    ]
                }
            
            elif command == "sora":
                # Generate video with Sora (via compatible model)
                st.info("ðŸŽ¨ Generating reference image for Sora...")
                image_url = self.replicate.generate_image(prompt=prompt, width=1280, height=720)
                st.info("ðŸŽ¬ Generating Sora-style video (this may take a few minutes)...")
                
                # Try text-to-video models
                try:
                    video_url = self.replicate._run_model(
                        "lightricks/ltx-video",
                        {"prompt": prompt}
                    )
                except Exception:
                    # Fallback to standard video
                    video_url = self.replicate.generate_video(
                        prompt=prompt,
                        image_url=image_url
                    )
                
                video_url = video_url if isinstance(video_url, str) else video_url[0] if video_url else None
                
                # Save to campaigns folder
                video_filename = f"sora_{timestamp}.mp4"
                image_filename = f"sora_ref_{timestamp}.png"
                video_path = self._save_to_campaigns(video_filename, url=video_url)
                image_path = self._save_to_campaigns(image_filename, url=image_url)
                
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "video",
                    "url": video_url,
                    "source_image": image_url,
                    "filepath": video_path,
                    "message": f"âœ… Generated video",
                    "artifacts": [
                        {"type": "video", "url": video_url, "filepath": video_path},
                        {"type": "image", "url": image_url, "filepath": image_path}
                    ]
                }
            
            elif command == "music":
                # Use MusicGen
                st.info("ðŸŽµ Generating music...")
                audio_url = self.replicate._run_model(
                    "meta/musicgen:large",
                    {
                        "prompt": prompt,
                        "duration": 30,
                        "model_version": "stereo-large"
                    }
                )
                
                # Save to campaigns folder
                audio_filename = f"music_{timestamp}.mp3"
                audio_path = self._save_to_campaigns(audio_filename, url=audio_url)
                
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "audio",
                    "url": audio_url,
                    "filepath": audio_path,
                    "message": f"âœ… Generated music",
                    "artifacts": [{"type": "audio", "url": audio_url, "filepath": audio_path, "prompt": prompt}]
                }
            
            elif command == "speak":
                # Use speech model
                st.info("ðŸ—£ï¸ Generating speech...")
                audio_url = self.replicate._run_model(
                    "minimax/speech-02-hd",
                    {"text": prompt}
                )
                
                # Save to campaigns folder
                audio_filename = f"speech_{timestamp}.mp3"
                audio_path = self._save_to_campaigns(audio_filename, url=audio_url)
                
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "audio",
                    "url": audio_url,
                    "filepath": audio_path,
                    "message": f"âœ… Generated speech audio",
                    "artifacts": [{"type": "audio", "url": audio_url, "filepath": audio_path, "text": prompt}]
                }
            
            elif command == "sound":
                # Sound effect generation
                st.info("ðŸ”Š Generating sound effect...")
                audio_url = self.replicate._run_model(
                    "stability-ai/stable-audio-2.5",
                    {"prompt": f"sound effect: {prompt}", "seconds_total": 10}
                )
                
                # Save to campaigns folder
                audio_filename = f"sound_{timestamp}.mp3"
                audio_path = self._save_to_campaigns(audio_filename, url=audio_url)
                
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "audio",
                    "url": audio_url,
                    "filepath": audio_path,
                    "message": f"âœ… Generated sound effect",
                    "artifacts": [{"type": "audio", "url": audio_url, "filepath": audio_path}]
                }
            
            elif command in ("3d", "model3d"):
                # 3D model generation - first generate an image, then convert to 3D
                st.info("ðŸŽ¨ Generating reference image for 3D model...")
                image_url = self.replicate.generate_image(prompt=prompt, width=1024, height=1024)
                st.info("ðŸŽ® Converting to 3D model...")
                
                # Use a model that accepts image input
                try:
                    model_url = self.replicate._run_model(
                        "stability-ai/stable-point-aware-3d",
                        {"image": image_url}
                    )
                except Exception:
                    # Fallback to text-based 3D
                    model_url = self.replicate._run_model(
                        "cjwbw/shap-e",
                        {"prompt": prompt}
                    )
                
                model_url = model_url if isinstance(model_url, str) else model_url[0] if model_url else None
                
                # Save to campaigns folder
                model_filename = f"model3d_{timestamp}.glb"
                image_filename = f"model3d_ref_{timestamp}.png"
                model_path = self._save_to_campaigns(model_filename, url=model_url)
                image_path = self._save_to_campaigns(image_filename, url=image_url)
                
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "3d",
                    "url": model_url or image_url,
                    "source_image": image_url,
                    "filepath": model_path,
                    "message": f"âœ… Generated 3D model",
                    "artifacts": [
                        {"type": "3d", "url": model_url, "filepath": model_path},
                        {"type": "image", "url": image_url, "filepath": image_path}
                    ]
                }
            
            return {"success": False, "error": f"Unknown media type: {command}"}
            
        except Exception as e:
            logger.error(f"Media generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    async def _call_model(self, model_name: str, prompt: str) -> Dict[str, Any]:
        """Call a specific AI model directly."""
        if model_name not in AI_MODELS:
            return {"success": False, "error": f"Unknown model: {model_name}"}
        
        model_id = AI_MODELS[model_name]
        
        try:
            # Determine model type and call appropriately
            model_lower = model_name.lower()
            
            # Image models
            if any(x in model_lower for x in ['flux', 'sdxl', 'imagen', 'bria', 'ideogram', 'seedream']):
                result = self.replicate._run_model(model_id, {"prompt": prompt})
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "image",
                    "url": result if isinstance(result, str) else result[0],
                    "model": model_id,
                    "message": f"âœ… Generated image with {model_name}",
                    "artifacts": [{"type": "image", "url": result if isinstance(result, str) else result[0], "model": model_id}]
                }
            
            # Music models
            elif any(x in model_lower for x in ['music', 'audio', 'lyria']):
                result = self.replicate._run_model(model_id, {"prompt": prompt, "duration": 30})
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "audio",
                    "url": result,
                    "model": model_id,
                    "message": f"âœ… Generated audio with {model_name}",
                    "artifacts": [{"type": "audio", "url": result, "model": model_id}]
                }
            
            # Speech models
            elif 'speech' in model_lower:
                result = self.replicate._run_model(model_id, {"text": prompt})
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "audio",
                    "url": result,
                    "model": model_id,
                    "message": f"âœ… Generated speech with {model_name}",
                    "artifacts": [{"type": "audio", "url": result, "model": model_id}]
                }
            
            # 3D models
            elif any(x in model_lower for x in ['3d', 'hunyuan', 'rodin', 'morphix', 'lucid', 'vggt']):
                result = self.replicate._run_model(model_id, {"prompt": prompt})
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "3d",
                    "url": result,
                    "model": model_id,
                    "message": f"âœ… Generated 3D with {model_name}",
                    "artifacts": [{"type": "3d", "url": result, "model": model_id}]
                }
            
            # Video models
            elif any(x in model_lower for x in ['video', 'kling', 'luma', 'minimax-video', 'sora', 'wan', 'ltx', 'mochi']):
                # Generate image first for video
                st.info(f"ðŸŽ¨ Generating reference image for video...")
                image_url = self.replicate.generate_image(prompt=prompt, width=1280, height=720)
                st.info(f"ðŸŽ¬ Generating video with {model_name}...")
                
                # Use the standard video generation method which handles model-specific params
                try:
                    result = self.replicate.generate_video(
                        prompt=prompt,
                        image_url=image_url
                    )
                except Exception as e:
                    # Fallback: try direct model call
                    result = self.replicate._run_model(model_id, {"prompt": prompt})
                
                video_url = result if isinstance(result, str) else result[0] if result else None
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "video",
                    "url": video_url,
                    "source_image": image_url,
                    "model": model_id,
                    "message": f"âœ… Generated video with {model_name}",
                    "artifacts": [
                        {"type": "video", "url": video_url},
                        {"type": "image", "url": image_url}
                    ]
                }
            
            # Text models
            elif any(x in model_lower for x in ['gpt', 'text', 'llm']):
                result = self.replicate._run_model(model_id, {"prompt": prompt})
                return {
                    "success": True,
                    "type": "text",
                    "content": result,
                    "model": model_id,
                    "message": f"âœ… Generated text with {model_name}",
                    "artifacts": [{"type": "text", "content": result, "model": model_id}]
                }
            
            # Ads/marketing models
            elif any(x in model_lower for x in ['ads', 'logo', 'inpaint', 'product']):
                result = self.replicate._run_model(model_id, {"prompt": prompt})
                return {
                    "success": True,
                    "type": "media",
                    "media_type": "image",
                    "url": result if isinstance(result, str) else result[0],
                    "model": model_id,
                    "message": f"âœ… Generated ad/marketing content with {model_name}",
                    "artifacts": [{"type": "image", "url": result if isinstance(result, str) else result[0], "model": model_id}]
                }
            
            # Default - try as image model
            else:
                result = self.replicate._run_model(model_id, {"prompt": prompt})
                return {
                    "success": True,
                    "type": "result",
                    "output": result,
                    "model": model_id,
                    "message": f"âœ… Model {model_name} completed",
                    "artifacts": [{"type": "output", "data": result, "model": model_id}]
                }
                
        except Exception as e:
            logger.error(f"Model call failed for {model_name}: {e}")
            return {"success": False, "error": str(e), "model": model_id}
    
    async def _execute_chain(self, chain_prompt: str, progress_callback: Callable = None) -> Dict[str, Any]:
        """
        Execute a chain of commands separated by | or ->
        Example: /chain /image cat | /video | /music relaxing
        Example: /chain /image sunset -> /video animate this -> /music chill vibes
        """
        # Support both | and -> as separators
        if '->' in chain_prompt:
            steps = [s.strip() for s in chain_prompt.split('->')]
        else:
            steps = [s.strip() for s in chain_prompt.split('|')]
        
        results = []
        context = {}  # Pass context between steps
        
        for i, step in enumerate(steps):
            step_name = step[:50] + "..." if len(step) > 50 else step
            
            if progress_callback:
                progress_callback({
                    "step": i + 1,
                    "total": len(steps),
                    "name": f"Chain Step {i+1}: {step_name}",
                    "status": "running"
                })
            
            # If step doesn't start with /, treat it as continuation of previous
            if not step.startswith('/'):
                # Use previous result's image/video for this step
                if context.get('last_image'):
                    step = f"/video {step}" if 'video' not in context.get('last_type', '') else f"/image {step}"
            
            # Execute the step
            result = await self.execute(step, progress_callback)
            results.append(result)
            
            # Update context for chaining
            if result.get('success'):
                if result.get('artifacts'):
                    for artifact in result['artifacts']:
                        if artifact.get('type') == 'image':
                            context['last_image'] = artifact.get('url')
                            context['last_type'] = 'image'
                        elif artifact.get('type') == 'video':
                            context['last_video'] = artifact.get('url')
                            context['last_type'] = 'video'
                        elif artifact.get('type') == 'audio':
                            context['last_audio'] = artifact.get('url')
                            context['last_type'] = 'audio'
            
            if progress_callback:
                progress_callback({
                    "step": i + 1,
                    "total": len(steps),
                    "name": f"Chain Step {i+1}: {step_name}",
                    "status": "completed" if result.get('success') else "failed",
                    "result": result
                })
        
        # Collect all artifacts
        all_artifacts = []
        for r in results:
            if r.get('artifacts'):
                all_artifacts.extend(r['artifacts'])
        
        success_count = sum(1 for r in results if r.get('success'))
        
        return {
            "success": success_count == len(results),
            "type": "chain",
            "steps_completed": success_count,
            "steps_total": len(results),
            "results": results,
            "message": f"âœ… Chain completed: {success_count}/{len(results)} steps successful",
            "artifacts": all_artifacts
        }


# Global slash command processor (initialized lazily)
_slash_processor: Optional[SlashCommandProcessor] = None

def get_slash_processor(replicate_api) -> SlashCommandProcessor:
    """Get or create the global slash command processor."""
    global _slash_processor
    if _slash_processor is None:
        _slash_processor = SlashCommandProcessor(replicate_api)
    return _slash_processor


class TaskType(Enum):
    """Types of tasks Otto can handle."""
    DESIGN = "design"
    VIDEO = "video"
    CONTENT = "content"
    SOCIAL = "social"
    PUBLISH = "publish"
    RESEARCH = "research"
    BROWSER = "browser"
    ANALYTICS = "analytics"
    CAMPAIGN = "campaign"
    GENERAL = "general"


@dataclass
class TaskStep:
    """A single step in a task workflow."""
    id: str
    name: str
    description: str
    task_type: TaskType
    status: str = "pending"  # pending, running, completed, failed
    result: Any = None
    error: str = None
    started_at: datetime = None
    completed_at: datetime = None
    artifacts: List[Dict] = field(default_factory=list)  # Generated files, URLs, etc.


@dataclass
class TaskPlan:
    """A complete plan for executing a user request."""
    id: str
    user_request: str
    summary: str
    steps: List[TaskStep]
    context: Dict = field(default_factory=dict)
    status: str = "planned"  # planned, running, completed, failed, paused
    created_at: datetime = field(default_factory=datetime.now)
    started_at: datetime = None
    completed_at: datetime = None
    final_result: str = None


class OttoEngine:
    """
    The core AI engine for Otto Mate - ENHANCED EDITION.
    Handles complex request understanding, planning, and execution.
    
    Features:
    - Intelligent request analysis with caching
    - Parallel step execution for speed
    - Smart retry with exponential backoff
    - Context-aware task routing
    """
    
    def __init__(self, replicate_api, printify_api=None, shopify_api=None, youtube_api=None):
        """Initialize Otto Engine with all service APIs."""
        self.replicate = replicate_api
        self.printify = printify_api
        self.shopify = shopify_api
        self.youtube = youtube_api
        self.executor = ThreadPoolExecutor(max_workers=OTTO_CONFIG.max_concurrent_tasks)
        
        # Browser automation (lazy loaded)
        self._browser_service = None
        self._anthropic_key = os.getenv('ANTHROPIC_API_KEY')
        
        # Current execution state
        self.current_plan: Optional[TaskPlan] = None
        self.execution_callbacks: List[Callable] = []
        
        # Performance tracking
        self._metrics: Dict[str, List[float]] = {
            'analysis_time': [],
            'execution_time': [],
            'api_calls': []
        }
        
        logger.info("ðŸ§  Otto Engine ENHANCED initialized")
    
    @property
    def browser_available(self) -> bool:
        """Check if browser automation is available."""
        return bool(self._anthropic_key)
    
    async def get_browser_service(self):
        """Lazy load browser service."""
        if self._browser_service is None:
            try:
                from automation.browser_service import BrowserAutomationService
                self._browser_service = BrowserAutomationService(self.replicate)
            except ImportError as e:
                logger.warning(f"Browser service not available: {e}")
        return self._browser_service
    
    def analyze_request(self, user_request: str) -> Dict[str, Any]:
        """
        Use AI to deeply analyze the user's request with ENHANCED intelligence.
        
        Features:
        - Request caching for repeated queries
        - Smarter prompt engineering
        - Better dependency detection
        - Cost/time estimation
        - Knowledge Base context injection
        - App state awareness
        """
        # Check cache first
        cached = _request_cache.get(user_request)
        if cached:
            return cached
        
        start_time = time.time()
        
        # Get knowledge base context
        kb = get_knowledge_base()
        kb_context = kb.get_context_summary()
        
        # Get app state
        app_state = OttoAppAwareness.get_state_summary()
        
        analysis_prompt = f"""You are Otto, an expert AI assistant for a marketing automation platform.
Analyze this request and create an OPTIMAL execution plan that minimizes time and maximizes quality.

USER REQUEST: {user_request}

KNOWLEDGE BASE CONTEXT (information the user has taught me):
{kb_context}

CURRENT APP STATE:
{app_state}

AVAILABLE CAPABILITIES (use the most efficient combination):
1. IMAGE GENERATION - Create designs, thumbnails, artwork (Flux Fast: $0.003, 10s | SDXL: $0.01, 20s)
2. VIDEO GENERATION - Create promotional videos (Kling: $0.20, 2min | Sora: $0.50, 3min)
3. TEXT GENERATION - Write descriptions, scripts, social posts (Fast, ~$0.001)
4. PRINTIFY SYNC - Upload designs, create/publish products
5. SHOPIFY SYNC - Create products, publish blogs, manage store
6. YOUTUBE SYNC - Upload videos with metadata
7. BROWSER AUTOMATION - Navigate websites, post to social media
8. ANALYTICS - Get shop statistics, orders, revenue
9. AI ASSISTANTS - Brand builder, outreach, campaign, design, content

OPTIMIZATION RULES:
- Steps with NO dependencies can run in PARALLEL (mark "parallel_group" same number)
- Use the fastest model that meets quality requirements
- Batch similar operations together
- Minimize API calls
- Use knowledge from the knowledge base when relevant

RESPOND IN THIS EXACT JSON FORMAT:
{{
    "understood_goal": "Brief summary of what user wants",
    "complexity": "simple|medium|complex",
    "requires_confirmation": false,
    "total_estimated_cost": "$X.XX",
    "total_estimated_time": "X minutes",
    "steps": [
        {{
            "step_number": 1,
            "name": "Step Name",
            "description": "What this step does",
            "task_type": "design|video|content|social|publish|research|browser|analytics|campaign",
            "service": "replicate|printify|shopify|youtube|browser|none",
            "depends_on": [],
            "parallel_group": 1,
            "estimated_time": "30 seconds",
            "estimated_cost": "$0.01"
        }}
    ],
    "expected_outputs": ["List of what will be produced"],
    "potential_issues": ["Any potential problems"]
}}

Be specific, actionable, and EFFICIENT."""

        try:
            response = self.replicate.generate_text(
                prompt=analysis_prompt,
                max_tokens=OTTO_CONFIG.max_tokens_analysis,
                temperature=OTTO_CONFIG.default_temperature
            )
            
            # Parse JSON from response
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                result = json.loads(json_match.group())
                
                # Cache the result
                _request_cache.set(user_request, result)
                
                # Track metrics
                elapsed = time.time() - start_time
                self._metrics['analysis_time'].append(elapsed)
                logger.info(f"ðŸ“Š Analysis completed in {elapsed:.2f}s")
                
                return result
            
        except Exception as e:
            logger.error(f"Request analysis failed: {e}")
        
        # Fallback analysis
        return self._fallback_analysis(user_request)
    
    def _fallback_analysis(self, request: str) -> Dict[str, Any]:
        """Simple pattern-based analysis as fallback."""
        request_lower = request.lower()
        steps = []
        
        # Detect task types from keywords
        if any(w in request_lower for w in ['design', 'image', 'create', 'make', 'generate']):
            if any(w in request_lower for w in ['product', 't-shirt', 'hoodie', 'mug', 'poster']):
                steps.append({
                    "step_number": 1,
                    "name": "Generate Product Design",
                    "description": f"Create design based on: {request}",
                    "task_type": "design",
                    "service": "replicate",
                    "depends_on": [],
                    "estimated_time": "30 seconds"
                })
        
        if any(w in request_lower for w in ['video', 'commercial', 'promo', 'ad']):
            steps.append({
                "step_number": len(steps) + 1,
                "name": "Generate Video",
                "description": "Create promotional video",
                "task_type": "video",
                "service": "replicate",
                "depends_on": [1] if steps else [],
                "estimated_time": "2 minutes"
            })
        
        if any(w in request_lower for w in ['printify', 'print', 'mockup']):
            steps.append({
                "step_number": len(steps) + 1,
                "name": "Upload to Printify",
                "description": "Create product on Printify",
                "task_type": "publish",
                "service": "printify",
                "depends_on": [1] if steps else [],
                "estimated_time": "30 seconds"
            })
        
        if any(w in request_lower for w in ['youtube', 'upload video']):
            steps.append({
                "step_number": len(steps) + 1,
                "name": "Upload to YouTube",
                "description": "Publish video to YouTube",
                "task_type": "publish",
                "service": "youtube",
                "depends_on": [],
                "estimated_time": "1 minute"
            })
        
        if any(w in request_lower for w in ['twitter', 'tweet', 'post', 'share']):
            steps.append({
                "step_number": len(steps) + 1,
                "name": "Post to Social Media",
                "description": "Share on social platforms",
                "task_type": "social",
                "service": "browser",
                "depends_on": [],
                "estimated_time": "1 minute"
            })
        
        # Default step if nothing detected
        if not steps:
            steps.append({
                "step_number": 1,
                "name": "Process Request",
                "description": request,
                "task_type": "general",
                "service": "replicate",
                "depends_on": [],
                "estimated_time": "30 seconds"
            })
        
        return {
            "understood_goal": request,
            "complexity": "medium" if len(steps) > 2 else "simple",
            "requires_confirmation": len(steps) > 3,
            "steps": steps,
            "expected_outputs": ["Generated content based on request"],
            "potential_issues": []
        }
    
    def create_plan(self, user_request: str) -> TaskPlan:
        """Create an execution plan from a user request."""
        analysis = self.analyze_request(user_request)
        
        steps = []
        for step_data in analysis.get("steps", []):
            step = TaskStep(
                id=str(uuid.uuid4())[:8],
                name=step_data.get("name", "Step"),
                description=step_data.get("description", ""),
                task_type=TaskType(step_data.get("task_type", "general"))
            )
            steps.append(step)
        
        plan = TaskPlan(
            id=str(uuid.uuid4())[:8],
            user_request=user_request,
            summary=analysis.get("understood_goal", user_request),
            steps=steps,
            context={
                "complexity": analysis.get("complexity", "medium"),
                "expected_outputs": analysis.get("expected_outputs", []),
                "original_analysis": analysis
            }
        )
        
        return plan
    
    async def execute_plan(self, plan: TaskPlan, progress_callback: Callable = None) -> TaskPlan:
        """
        Execute a task plan with PARALLEL execution for independent steps.
        
        ENHANCED Features:
        - Parallel execution of independent steps
        - Smart retry with exponential backoff
        - Real-time progress updates
        - Performance metrics tracking
        
        Args:
            plan: The TaskPlan to execute
            progress_callback: Optional callback for progress updates
        """
        self.current_plan = plan
        plan.status = "running"
        plan.started_at = datetime.now()
        start_time = time.time()
        
        total_steps = len(plan.steps)
        
        # Group steps by parallel_group for concurrent execution
        if OTTO_CONFIG.enable_parallel_steps:
            step_groups = self._group_parallel_steps(plan.steps)
        else:
            step_groups = [[step] for step in plan.steps]
        
        completed_count = 0
        
        for group_idx, step_group in enumerate(step_groups):
            if len(step_group) > 1:
                # Execute steps in parallel
                logger.info(f"âš¡ Executing {len(step_group)} steps in parallel")
                tasks = []
                for step in step_group:
                    step.status = "running"
                    step.started_at = datetime.now()
                    tasks.append(self._execute_step_with_retry(step, plan.context))
                
                # Wait for all parallel steps
                results = await asyncio.gather(*tasks, return_exceptions=True)
                
                for step, result in zip(step_group, results):
                    completed_count += 1
                    if isinstance(result, Exception):
                        step.status = "failed"
                        step.error = str(result)
                        logger.error(f"Step {step.name} failed: {result}")
                    else:
                        step.result = result
                        step.status = "completed"
                        if result.get("artifacts"):
                            step.artifacts = result["artifacts"]
                            plan.context["latest_artifacts"] = result["artifacts"]
                    
                    step.completed_at = datetime.now()
                    
                    if progress_callback:
                        progress_callback({
                            "step": completed_count,
                            "total": total_steps,
                            "name": step.name,
                            "status": step.status,
                            "result": step.result if step.status == "completed" else None,
                            "error": step.error if step.status == "failed" else None
                        })
            else:
                # Execute single step
                step = step_group[0]
                step.status = "running"
                step.started_at = datetime.now()
                completed_count += 1
                
                if progress_callback:
                    progress_callback({
                        "step": completed_count,
                        "total": total_steps,
                        "name": step.name,
                        "status": "running"
                    })
                
                try:
                    result = await self._execute_step_with_retry(step, plan.context)
                    
                    step.result = result
                    step.status = "completed"
                    step.completed_at = datetime.now()
                    
                    plan.context[f"step_{completed_count}_result"] = result
                    if result.get("artifacts"):
                        step.artifacts = result["artifacts"]
                        plan.context["latest_artifacts"] = result["artifacts"]
                    
                    if progress_callback:
                        progress_callback({
                            "step": completed_count,
                            "total": total_steps,
                            "name": step.name,
                            "status": "completed",
                            "result": result
                        })
                    
                except Exception as e:
                    logger.error(f"Step {step.name} failed: {e}")
                    step.status = "failed"
                    step.error = str(e)
                    step.completed_at = datetime.now()
                    
                    if progress_callback:
                        progress_callback({
                            "step": completed_count,
                            "total": total_steps,
                            "name": step.name,
                            "status": "failed",
                            "error": str(e)
                        })
                    
                    # Continue with other steps unless critical
                    if step.task_type in [TaskType.DESIGN]:
                        plan.status = "failed"
                        break
        
        # Mark plan as completed
        if plan.status != "failed":
            plan.status = "completed"
        plan.completed_at = datetime.now()
        
        # Track execution time
        elapsed = time.time() - start_time
        self._metrics['execution_time'].append(elapsed)
        logger.info(f"ðŸ“Š Plan executed in {elapsed:.2f}s")
        
        # Generate final summary
        plan.final_result = self._generate_summary(plan)
        
        return plan
    
    def _group_parallel_steps(self, steps: List[TaskStep]) -> List[List[TaskStep]]:
        """Group steps that can run in parallel based on dependencies."""
        groups = []
        current_group = []
        last_group_num = None
        
        for step in steps:
            # Check for parallel_group in the step context
            group_num = getattr(step, 'parallel_group', None)
            
            if group_num is not None and group_num == last_group_num:
                current_group.append(step)
            else:
                if current_group:
                    groups.append(current_group)
                current_group = [step]
                last_group_num = group_num
        
        if current_group:
            groups.append(current_group)
        
        return groups if groups else [[step] for step in steps]
    
    async def _execute_step_with_retry(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Execute a step with exponential backoff retry."""
        last_error = None
        
        for attempt in range(OTTO_CONFIG.max_retries):
            try:
                return await self._execute_step(step, context)
            except Exception as e:
                last_error = e
                if attempt < OTTO_CONFIG.max_retries - 1:
                    delay = min(
                        OTTO_CONFIG.retry_base_delay * (2 ** attempt),
                        OTTO_CONFIG.retry_max_delay
                    )
                    logger.warning(f"Step {step.name} failed (attempt {attempt + 1}), retrying in {delay}s: {e}")
                    await asyncio.sleep(delay)
        
        raise last_error

    async def _execute_step(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Execute a single step and return results."""
        
        if step.task_type == TaskType.DESIGN:
            return await self._execute_design(step, context)
        
        elif step.task_type == TaskType.VIDEO:
            return await self._execute_video(step, context)
        
        elif step.task_type == TaskType.CONTENT:
            return await self._execute_content(step, context)
        
        elif step.task_type == TaskType.SOCIAL:
            return await self._execute_social(step, context)
        
        elif step.task_type == TaskType.PUBLISH:
            return await self._execute_publish(step, context)
        
        elif step.task_type == TaskType.BROWSER:
            return await self._execute_browser(step, context)
        
        elif step.task_type == TaskType.ANALYTICS:
            return await self._execute_analytics(step, context)
        
        elif step.task_type == TaskType.CAMPAIGN:
            return await self._execute_campaign(step, context)
        
        else:
            return await self._execute_general(step, context)
    
    async def _execute_design(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Generate images/designs."""
        prompt = step.description
        
        # Enhance prompt for better results
        enhanced_prompt = f"Professional product design: {prompt}. High quality, commercial ready, clean background, centered composition, vibrant colors."
        
        try:
            image_url = self.replicate.generate_image(
                prompt=enhanced_prompt,
                width=1024,
                height=1024,
                aspect_ratio="1:1"
            )
            
            return {
                "type": "image",
                "url": image_url,
                "prompt": enhanced_prompt,
                "artifacts": [{"type": "image", "url": image_url, "prompt": prompt}]
            }
        except Exception as e:
            raise Exception(f"Image generation failed: {e}")
    
    async def _execute_video(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Generate videos from images or prompts."""
        # Check for image from previous step
        image_url = None
        if "latest_artifacts" in context:
            for artifact in context["latest_artifacts"]:
                if artifact.get("type") == "image":
                    image_url = artifact.get("url")
                    break
        
        try:
            video_prompt = f"Gentle motion, professional showcase: {step.description}"
            
            if image_url:
                video_url = self.replicate.generate_video(
                    prompt=video_prompt,
                    image_url=image_url,
                    aspect_ratio="16:9",
                    motion_level=3
                )
            else:
                # Generate image first, then video
                image_url = self.replicate.generate_image(
                    prompt=step.description,
                    width=1280,
                    height=720,
                    aspect_ratio="16:9"
                )
                video_url = self.replicate.generate_video(
                    prompt=video_prompt,
                    image_url=image_url,
                    aspect_ratio="16:9",
                    motion_level=3
                )
            
            return {
                "type": "video",
                "url": video_url,
                "source_image": image_url,
                "artifacts": [
                    {"type": "video", "url": video_url},
                    {"type": "image", "url": image_url}
                ]
            }
        except Exception as e:
            raise Exception(f"Video generation failed: {e}")
    
    async def _execute_content(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Generate text content (descriptions, scripts, posts)."""
        content_prompt = f"""Generate professional marketing content for: {step.description}

Context from previous steps: {json.dumps(context.get('latest_artifacts', [])[:2])}

Provide:
1. A compelling product description (2 paragraphs)
2. 3 social media posts (Twitter, Instagram, Facebook)
3. 5 relevant hashtags

Format with clear headers."""

        try:
            content = self.replicate.generate_text(
                prompt=content_prompt,
                max_tokens=800,
                temperature=0.7
            )
            
            return {
                "type": "content",
                "text": content,
                "artifacts": [{"type": "text", "content": content}]
            }
        except Exception as e:
            raise Exception(f"Content generation failed: {e}")
    
    async def _execute_social(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Post to social media using browser automation."""
        if not self.browser_available:
            return {
                "type": "social",
                "status": "skipped",
                "message": "Browser automation not available (ANTHROPIC_API_KEY not set)",
                "artifacts": []
            }
        
        try:
            browser = await self.get_browser_service()
            if browser:
                result = await browser.execute(
                    task=step.description,
                    max_steps=15
                )
                
                return {
                    "type": "social",
                    "status": "completed" if result.get("success") else "failed",
                    "result": result.get("result"),
                    "artifacts": []
                }
            else:
                return {
                    "type": "social",
                    "status": "skipped",
                    "message": "Browser service unavailable",
                    "artifacts": []
                }
        except Exception as e:
            raise Exception(f"Social posting failed: {e}")
    
    async def _execute_publish(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Publish to services (Printify, Shopify, YouTube)."""
        description_lower = step.description.lower()
        
        # Determine which service
        if "printify" in description_lower or "print" in description_lower:
            return await self._publish_to_printify(step, context)
        
        elif "youtube" in description_lower:
            return await self._publish_to_youtube(step, context)
        
        elif "shopify" in description_lower:
            return await self._publish_to_shopify(step, context)
        
        return {
            "type": "publish",
            "status": "skipped",
            "message": "Unknown publish target",
            "artifacts": []
        }
    
    async def _publish_to_printify(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Upload design to Printify and create product."""
        if not self.printify:
            return {
                "type": "publish",
                "service": "printify",
                "status": "skipped",
                "message": "âŒ Printify API not configured. Add your API key in Settings â†’ Integrations.",
                "artifacts": []
            }
        
        # Get image from context
        image_url = None
        if "latest_artifacts" in context:
            for artifact in context["latest_artifacts"]:
                if artifact.get("type") == "image":
                    image_url = artifact.get("url")
                    break
        
        # Also check first generated image
        if not image_url and context.get("step_1_result"):
            first_result = context["step_1_result"]
            if first_result.get("type") == "image":
                image_url = first_result.get("url")
        
        if not image_url:
            return {
                "type": "publish",
                "service": "printify",
                "status": "failed",
                "message": "No image available to upload",
                "artifacts": []
            }
        
        try:
            import requests
            
            # Download image as bytes
            logger.info(f"ðŸ“¥ Downloading image from: {image_url[:50]}...")
            response = requests.get(image_url, timeout=60)
            response.raise_for_status()
            image_bytes = response.content
            
            # Upload to Printify - correct signature: upload_image(image_data: bytes, file_name: str)
            file_name = f"otto_design_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            logger.info(f"ðŸ“¤ Uploading to Printify as: {file_name}")
            
            printify_image_id = self.printify.upload_image(image_bytes, file_name)
            
            logger.info(f"âœ… Printify upload successful! Image ID: {printify_image_id}")
            
            # Get shop ID
            shop_id = None
            try:
                shops = self.printify.get_shops()
                if shops:
                    shop_id = shops[0].get('id')
            except Exception:
                pass
            
            return {
                "type": "publish",
                "service": "printify",
                "status": "completed",
                "printify_image_id": printify_image_id,
                "shop_id": shop_id,
                "message": f"âœ… Image uploaded to Printify!\n\n**Image ID:** `{printify_image_id}`\n\nYou can now create a product with this design in your Printify dashboard.",
                "artifacts": [{"type": "printify_upload", "id": printify_image_id, "shop_id": shop_id}]
            }
        except Exception as e:
            logger.error(f"Printify upload failed: {e}")
            raise Exception(f"Printify upload failed: {e}")
    
    async def _publish_to_youtube(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Upload video to YouTube."""
        if not self.youtube:
            return {
                "type": "publish",
                "service": "youtube",
                "status": "skipped",
                "message": "âŒ YouTube not configured. Set up YouTube OAuth in Settings â†’ Integrations.",
                "artifacts": []
            }
        
        # Get video from context - check multiple places
        video_url = None
        
        # Check latest_artifacts
        if "latest_artifacts" in context:
            for artifact in context["latest_artifacts"]:
                if artifact.get("type") == "video":
                    video_url = artifact.get("url")
                    break
        
        # Check all step results for videos
        if not video_url:
            for key, value in context.items():
                if key.startswith("step_") and isinstance(value, dict):
                    if value.get("type") == "video":
                        video_url = value.get("url")
                        break
        
        if not video_url:
            return {
                "type": "publish",
                "service": "youtube",
                "status": "failed",
                "message": "No video available to upload. Generate a video first.",
                "artifacts": []
            }
        
        try:
            import requests
            import tempfile
            
            logger.info(f"ðŸ“¥ Downloading video from: {video_url[:50]}...")
            response = requests.get(video_url, timeout=180)
            response.raise_for_status()
            
            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as f:
                f.write(response.content)
                video_path = f.name
            
            logger.info(f"ðŸ“¤ Uploading to YouTube...")
            
            # Get title from context
            title = context.get("understood_goal", "Otto Mate Creation")[:100]
            if len(title) < 10:
                title = f"Otto Mate - {step.description[:80]}"
            
            description = f"""Created with Otto Mate AI

{step.description}

#OttoMate #AIGenerated #Automated"""
            
            # Use upload_commercial method
            result = self.youtube.upload_commercial(
                video_path=video_path,
                product_name=title,
                metadata={
                    'title': title,
                    'description': description,
                    'tags': ['AI Generated', 'Otto Mate', 'Automated', 'Product Video'],
                    'privacy': 'unlisted'  # Unlisted so it's accessible but not public
                }
            )
            
            # Cleanup temp file
            try:
                os.unlink(video_path)
            except OSError:
                pass
            
            if result:
                video_id = result.get("id")
                video_link = result.get("url", f"https://youtube.com/watch?v={video_id}")
                
                logger.info(f"âœ… YouTube upload successful: {video_link}")
                
                return {
                    "type": "publish",
                    "service": "youtube",
                    "status": "completed",
                    "video_id": video_id,
                    "video_link": video_link,
                    "message": f"âœ… Video uploaded to YouTube!\n\n**URL:** {video_link}\n**Status:** Unlisted",
                    "artifacts": [{"type": "youtube_video", "id": video_id, "url": video_link}]
                }
            else:
                raise Exception("Upload returned no result")
                
        except Exception as e:
            logger.error(f"YouTube upload failed: {e}")
            raise Exception(f"YouTube upload failed: {e}")
    
    async def _publish_to_shopify(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Create blog post or product on Shopify."""
        if not self.shopify:
            return {
                "type": "publish",
                "service": "shopify",
                "status": "skipped",
                "message": "âŒ Shopify not configured. Add credentials in Settings â†’ Integrations.",
                "artifacts": []
            }
        
        description_lower = step.description.lower()
        
        # Determine if this is a blog post or product
        is_blog = any(word in description_lower for word in ['blog', 'article', 'post', 'write', 'seo'])
        
        # Get image from context
        image_url = None
        if "latest_artifacts" in context:
            for artifact in context["latest_artifacts"]:
                if artifact.get("type") == "image":
                    image_url = artifact.get("url")
                    break
        
        # Check step results for images
        if not image_url:
            for key, value in context.items():
                if key.startswith("step_") and isinstance(value, dict):
                    if value.get("type") == "image":
                        image_url = value.get("url")
                        break
        
        try:
            title = context.get("understood_goal", "Otto Mate Creation")[:100]
            
            if is_blog:
                # Generate SEO-friendly blog content
                logger.info("ðŸ“ Generating SEO blog content...")
                
                blog_prompt = f"""Write an SEO-optimized blog post for: {step.description}

Requirements:
- HTML formatted with proper tags (h2, h3, p, ul, li)
- Include meta description in first paragraph
- Use keywords naturally throughout
- 400-600 words
- Include a compelling introduction and conclusion
- Add relevant subheadings

Topic context: {context.get('understood_goal', step.description)}
"""
                
                blog_content = self.replicate.generate_text(
                    prompt=blog_prompt,
                    max_tokens=1200,
                    temperature=0.7
                )
                
                # Clean up the content
                if not blog_content.strip().startswith('<'):
                    blog_content = f"<p>{blog_content}</p>"
                
                # Add image to blog if available
                if image_url:
                    blog_content = f'<img src="{image_url}" alt="{title}" style="width:100%;max-width:800px;margin:20px 0;" />\n\n' + blog_content
                
                # Create the blog post
                logger.info(f"ðŸ“¤ Publishing blog to Shopify...")
                
                result = self.shopify.create_blog_post(
                    title=title,
                    body_html=blog_content,
                    author="Otto Mate AI",
                    tags=["AI Generated", "Otto Mate"],
                    published=True,
                    image_url=image_url
                )
                
                if result:
                    article_id = result.get('id')
                    article_url = result.get('url', f"Article ID: {article_id}")
                    
                    logger.info(f"âœ… Blog published: {article_url}")
                    
                    return {
                        "type": "publish",
                        "service": "shopify",
                        "content_type": "blog",
                        "status": "completed",
                        "article_id": article_id,
                        "article_url": article_url,
                        "message": f"âœ… Blog post published to Shopify!\n\n**URL:** {article_url}\n**ID:** {article_id}",
                        "artifacts": [{"type": "shopify_blog", "id": article_id, "url": article_url}]
                    }
                else:
                    raise Exception("Blog creation returned no result")
            
            else:
                # Create product
                product_data = {
                    "title": title,
                    "body_html": f"<p>{step.description}</p><p>Created with Otto Mate AI.</p>",
                    "vendor": "Otto Mate AI",
                    "product_type": "AI Generated",
                    "status": "draft"
                }
                
                if image_url:
                    product_data["images"] = [{"src": image_url}]
                
                result = self.shopify.create_product(product_data)
                product_id = result.get("product", {}).get("id")
                
                logger.info(f"âœ… Product created: {product_id}")
                
                return {
                    "type": "publish",
                    "service": "shopify",
                    "content_type": "product",
                    "status": "completed",
                    "product_id": product_id,
                    "message": f"âœ… Product created on Shopify (Draft)!\n\n**Product ID:** {product_id}",
                    "artifacts": [{"type": "shopify_product", "id": product_id}]
                }
                
        except Exception as e:
            logger.error(f"Shopify publish failed: {e}")
            raise Exception(f"Shopify publish failed: {e}")
    
    async def _execute_browser(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Execute browser automation tasks."""
        if not self.browser_available:
            return {
                "type": "browser",
                "status": "skipped",
                "message": "âŒ Browser automation not available. Set ANTHROPIC_API_KEY in .env file.",
                "artifacts": []
            }
        
        try:
            browser = await self.get_browser_service()
            if browser:
                logger.info(f"ðŸŒ Executing browser task: {step.description[:50]}...")
                
                result = await browser.execute(
                    task=step.description,
                    max_steps=20
                )
                
                if result.get("success"):
                    return {
                        "type": "browser",
                        "status": "completed",
                        "result": result.get("result", "Task completed"),
                        "steps_taken": result.get("steps", 0),
                        "message": f"âœ… Browser task completed in {result.get('steps', 0)} steps",
                        "artifacts": []
                    }
                else:
                    return {
                        "type": "browser",
                        "status": "failed",
                        "result": result.get("result"),
                        "errors": result.get("errors", []),
                        "message": f"âš ï¸ Browser task had issues: {result.get('errors', ['Unknown error'])}",
                        "artifacts": []
                    }
            else:
                return {
                    "type": "browser",
                    "status": "failed",
                    "message": "Browser service could not be initialized",
                    "artifacts": []
                }
        except Exception as e:
            logger.error(f"Browser automation failed: {e}")
            return {
                "type": "browser",
                "status": "failed",
                "message": f"Browser automation failed: {e}",
                "artifacts": []
            }
    
    async def _execute_analytics(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Get analytics from connected services."""
        results = {}
        
        if self.shopify:
            try:
                analytics = self.shopify.get_comprehensive_analytics()
                results["shopify"] = analytics
            except Exception as e:
                results["shopify_error"] = str(e)
        
        if self.printify:
            try:
                shops = self.printify.get_shops()
                results["printify_shops"] = shops
            except Exception as e:
                results["printify_error"] = str(e)
        
        return {
            "type": "analytics",
            "data": results,
            "artifacts": []
        }
    
    async def _execute_campaign(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """Execute full campaign generation."""
        # This delegates to the campaign generator
        return {
            "type": "campaign",
            "status": "delegated",
            "message": "Campaign generation should use Campaign Generator tab for full workflow",
            "artifacts": []
        }
    
    async def _execute_general(self, step: TaskStep, context: Dict) -> Dict[str, Any]:
        """General task execution using AI."""
        try:
            response = self.replicate.generate_text(
                prompt=step.description,
                max_tokens=600,
                temperature=0.7
            )
            
            return {
                "type": "text",
                "content": response,
                "artifacts": [{"type": "text", "content": response}]
            }
        except Exception as e:
            raise Exception(f"General execution failed: {e}")
    
    def _generate_summary(self, plan: TaskPlan) -> str:
        """Generate a human-readable summary of the execution."""
        completed = sum(1 for s in plan.steps if s.status == "completed")
        failed = sum(1 for s in plan.steps if s.status == "failed")
        
        summary_parts = [
            f"## ðŸŽ¯ Task Complete",
            f"",
            f"**Request:** {plan.summary}",
            f"**Status:** {'âœ… Success' if plan.status == 'completed' else 'âš ï¸ Completed with issues'}",
            f"**Steps:** {completed}/{len(plan.steps)} completed",
            f""
        ]
        
        # Add artifacts
        all_artifacts = []
        for step in plan.steps:
            all_artifacts.extend(step.artifacts)
        
        if all_artifacts:
            summary_parts.append("### ðŸ“¦ Generated Assets")
            for artifact in all_artifacts:
                if artifact.get("type") == "image":
                    summary_parts.append(f"ðŸ–¼ï¸ **Image:** {artifact.get('url', 'N/A')}")
                elif artifact.get("type") == "video":
                    summary_parts.append(f"ðŸŽ¬ **Video:** {artifact.get('url', 'N/A')}")
                elif artifact.get("type") == "text":
                    content = artifact.get("content", "")[:200]
                    summary_parts.append(f"ðŸ“ **Content:** {content}...")
                elif artifact.get("type") == "printify_upload":
                    summary_parts.append(f"ðŸ›ï¸ **Printify Upload:** ID {artifact.get('id')}")
                elif artifact.get("type") == "youtube_video":
                    summary_parts.append(f"ðŸ“º **YouTube:** {artifact.get('url')}")
                elif artifact.get("type") == "shopify_product":
                    summary_parts.append(f"ðŸª **Shopify Product:** ID {artifact.get('id')}")
        
        # Add failed steps
        if failed > 0:
            summary_parts.append("")
            summary_parts.append("### âš ï¸ Issues")
            for step in plan.steps:
                if step.status == "failed":
                    summary_parts.append(f"- {step.name}: {step.error}")
        
        return "\n".join(summary_parts)


def render_otto_chat(replicate_api, printify_api=None, shopify_api=None, youtube_api=None):
    """
    Render the Otto chat interface with real-time execution display.
    ENHANCED v3.0 with Knowledge Base and App Awareness
    """
    
    # Custom CSS for modern chat
    st.markdown("""
    <style>
    .otto-chat-container {
        display: flex;
        flex-direction: column;
        height: 100%;
    }
    .otto-progress {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
        color: white;
    }
    .otto-step {
        background: #f8f9fa;
        border-left: 4px solid #667eea;
        padding: 10px 15px;
        margin: 5px 0;
        border-radius: 0 8px 8px 0;
    }
    .otto-step-running {
        border-left-color: #ffc107;
        background: #fffbeb;
    }
    .otto-step-completed {
        border-left-color: #28a745;
        background: #f0fff4;
    }
    .otto-step-failed {
        border-left-color: #dc3545;
        background: #fff5f5;
    }
    .otto-artifact {
        background: white;
        border: 1px solid #e0e0e0;
        border-radius: 12px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    .otto-kb-upload {
        background: linear-gradient(135deg, #f5f7fa 0%, #e4e9f2 100%);
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
        border: 2px dashed #667eea;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Initialize engine
    engine = OttoEngine(
        replicate_api=replicate_api,
        printify_api=printify_api,
        shopify_api=shopify_api,
        youtube_api=youtube_api
    )
    
    # Initialize knowledge base
    kb = get_knowledge_base()
    
    # Initialize chat state
    if 'otto_messages' not in st.session_state:
        st.session_state.otto_messages = []
    if 'otto_current_plan' not in st.session_state:
        st.session_state.otto_current_plan = None
    
    # Header with status and knowledge base
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        st.markdown("### ðŸ¤– Otto Mate v3.0")
    with col2:
        services = []
        if printify_api:
            services.append("ðŸ›ï¸")
        if shopify_api:
            services.append("ðŸª")
        if youtube_api:
            services.append("ðŸ“º")
        if engine.browser_available:
            services.append("ðŸŒ")
        st.markdown(f"**Connected:** {' '.join(services) if services else 'âšª None'}")
    with col3:
        kb_stats = f"ðŸ“š {len(kb.facts)}F | {len(kb.documents)}D | {len(kb.images)}I"
        st.markdown(f"**Knowledge:** {kb_stats}")
    
    # Knowledge Base File Upload Section - Make more prominent
    st.markdown("---")
    with st.expander("ðŸ“š **Knowledge Base** - Teach Otto About Your Business", expanded=True):
        st.markdown("""
        **Upload images or documents to teach Otto context about your business, products, brand, or anything else.**
        Otto will analyze and remember this information for future conversations.
        """)
        
        upload_cols = st.columns(2)
        
        with upload_cols[0]:
            st.markdown("**ðŸ–¼ï¸ Image Analysis**")
            uploaded_image = st.file_uploader(
                "Upload an image to analyze",
                type=['png', 'jpg', 'jpeg', 'gif', 'webp'],
                key="otto_kb_image_upload"
            )
            
            if uploaded_image:
                if st.button("ðŸ” Analyze Image", key="analyze_image_btn"):
                    with st.spinner("Analyzing image..."):
                        image_data = uploaded_image.read()
                        result = kb.analyze_image(image_data, uploaded_image.name, replicate_api)
                        
                        if result.get("success"):
                            st.success(result.get("message"))
                            with st.expander("ðŸ“ Analysis Details"):
                                st.markdown(result.get("analysis", ""))
                            # Add to chat history
                            st.session_state.otto_messages.append({
                                'role': 'assistant',
                                'content': f"ðŸ“· **Image Analyzed:** {uploaded_image.name}\n\n{result.get('analysis', '')}",
                                'artifacts': []
                            })
                        else:
                            st.error(f"Analysis failed: {result.get('error')}")
        
        with upload_cols[1]:
            st.markdown("**ðŸ“„ Document Reading**")
            uploaded_doc = st.file_uploader(
                "Upload a document to read",
                type=['txt', 'pdf', 'doc', 'docx', 'md', 'json', 'csv'],
                key="otto_kb_doc_upload"
            )
            
            if uploaded_doc:
                if st.button("ðŸ“– Read Document", key="read_doc_btn"):
                    with st.spinner("Reading document..."):
                        doc_data = uploaded_doc.read()
                        result = kb.read_document(doc_data, uploaded_doc.name, replicate_api)
                        
                        if result.get("success"):
                            st.success(result.get("message"))
                            with st.expander("ðŸ“ Document Summary"):
                                st.markdown(result.get("summary", ""))
                            # Add to chat history
                            st.session_state.otto_messages.append({
                                'role': 'assistant',
                                'content': f"ðŸ“„ **Document Added:** {uploaded_doc.name}\n\n**Summary:**\n{result.get('summary', '')}",
                                'artifacts': []
                            })
                        else:
                            st.error(f"Reading failed: {result.get('error')}")
        
        # Quick fact input
        st.markdown("---")
        st.markdown("**ðŸ’¡ Quick Add Fact**")
        fact_cols = st.columns([4, 1])
        with fact_cols[0]:
            quick_fact = st.text_input(
                "Teach Otto something",
                placeholder="e.g., Our brand color is blue, we sell eco-friendly products...",
                key="otto_quick_fact_input",
                label_visibility="collapsed"
            )
        with fact_cols[1]:
            if st.button("âž• Add", key="add_fact_btn", use_container_width=True):
                if quick_fact:
                    kb.add_fact(quick_fact)
                    st.success(f"âœ… Learned: {quick_fact[:50]}...")
                    st.session_state.otto_messages.append({
                        'role': 'assistant',
                        'content': f"ðŸ“š **Learned new fact:** {quick_fact}",
                        'artifacts': []
                    })
                    st.rerun()
        
        # Quick knowledge actions
        st.markdown("---")
        kb_action_cols = st.columns(4)
        with kb_action_cols[0]:
            if st.button("ðŸ“Š Show Status", key="kb_status_btn", use_container_width=True):
                summary = kb.get_context_summary()
                st.info(summary)
        with kb_action_cols[1]:
            if st.button("ðŸ§¹ Clear Facts", key="kb_clear_facts_btn", use_container_width=True):
                kb.clear_memory("facts")
                st.success("Facts cleared!")
        with kb_action_cols[2]:
            if st.button("ðŸ—‘ï¸ Clear Docs", key="kb_clear_docs_btn", use_container_width=True):
                kb.clear_memory("documents")
                st.success("Documents cleared!")
        with kb_action_cols[3]:
            if st.button("ðŸ’¥ Clear All", key="kb_clear_all_btn", use_container_width=True):
                kb.clear_memory("all")
                st.success("All knowledge cleared!")
    
    # Messages container
    messages_container = st.container(height=400)
    
    with messages_container:
        if not st.session_state.otto_messages:
            st.markdown("""
            <div style="text-align: center; padding: 40px; color: #666;">
                <div style="font-size: 3em; margin-bottom: 15px;">ðŸ§ </div>
                <h3>Hey! I'm Otto Mate v3.0</h3>
                <p>I can create designs, videos, marketing campaigns, and publish to your connected services.</p>
                <p>Now with <strong>Knowledge Base</strong> - I remember everything you teach me!</p>
                <hr style="margin: 20px 0; border-color: #eee;">
                <p style="font-size: 0.9em;"><strong>âš¡ New AI Assistants:</strong></p>
                <p style="font-size: 0.85em; color: #888;">
                    <code>/brand</code> <code>/outreach</code> <code>/campaign</code> <code>/design</code> <code>/content</code>
                </p>
                <p style="font-size: 0.9em;"><strong>ðŸ“š Knowledge Commands:</strong></p>
                <p style="font-size: 0.85em; color: #888;">
                    <code>/learn</code> <code>/recall</code> <code>/knowledge</code> <code>/status</code>
                </p>
                <p style="font-size: 0.9em;"><strong>ðŸŽ¨ Media & Files:</strong></p>
                <p style="font-size: 0.85em; color: #888;">
                    <code>/image</code> <code>/video</code> <code>/music</code> <code>/python</code> <code>/pdf</code>
                </p>
                <p style="font-size: 0.8em; color: #999;">Type <code>/help</code> for all commands</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            for msg in st.session_state.otto_messages:
                with st.chat_message(msg['role'], avatar="ðŸ§‘" if msg['role'] == 'user' else "ðŸ¤–"):
                    st.markdown(msg['content'])
                    
                    # Render artifacts with proper media display
                    if msg.get('artifacts'):
                        for artifact in msg['artifacts']:
                            artifact_type = artifact.get('type', '')
                            artifact_url = artifact.get('url', '')
                            
                            if artifact_type == 'image':
                                st.image(artifact_url, use_container_width=True)
                            elif artifact_type == 'video':
                                st.video(artifact_url)
                            elif artifact_type == 'youtube_video':
                                # Embed YouTube video
                                video_id = artifact.get('id', '')
                                if video_id:
                                    st.markdown(f"""
                                    <iframe width="100%" height="400" 
                                        src="https://www.youtube.com/embed/{video_id}" 
                                        frameborder="0" allowfullscreen>
                                    </iframe>
                                    """, unsafe_allow_html=True)
                                elif artifact_url:
                                    st.video(artifact_url)
                            elif artifact_type == 'printify_upload':
                                st.success(f"âœ… Uploaded to Printify - Image ID: {artifact.get('id')}")
                            elif artifact_type == 'shopify_product':
                                st.success(f"âœ… Published to Shopify - Product ID: {artifact.get('id')}")
                                if artifact_url:
                                    st.markdown(f"[View Product]({artifact_url})")
                            elif artifact_type == 'shopify_blog':
                                st.success(f"âœ… Published to Shopify Blog - Article ID: {artifact.get('id')}")
                                if artifact_url:
                                    st.markdown(f"[View Blog Post]({artifact_url})")
                            elif artifact_type == 'text':
                                with st.expander("ðŸ“ Generated Content"):
                                    st.markdown(artifact.get('content', ''))
                            elif artifact_type == 'file':
                                filepath = artifact.get('path', artifact.get('filepath', ''))
                                filename = artifact.get('filename', 'download')
                                st.success(f"ðŸ“ Generated: **{filename}**")
                                if filepath and os.path.exists(filepath):
                                    with open(filepath, 'rb') as f:
                                        st.download_button(
                                            f"â¬‡ï¸ Download {filename}",
                                            f.read(),
                                            file_name=filename,
                                            key=f"download_{filename}_{id(artifact)}"
                                        )
                            elif artifact_type == 'audio':
                                st.audio(artifact_url)
    
    # Smart Command Suggestions
    st.markdown("---")
    
    # Inline autocomplete using session state
    if 'otto_input_buffer' not in st.session_state:
        st.session_state.otto_input_buffer = ""
    
    # Show suggestions when typing starts with /
    suggestion_cols = st.columns([4, 1])
    with suggestion_cols[1]:
        if st.button("â“ Commands", key="show_cmd_help", help="Show all slash commands"):
            st.session_state.show_command_palette = not st.session_state.get('show_command_palette', False)
    
    # Command palette
    if st.session_state.get('show_command_palette', False):
        with st.expander("âš¡ Quick Commands", expanded=True):
            st.markdown("**ðŸ§  AI Assistants**")
            cmd_cols = st.columns(5)
            ai_cmds = [
                ("ðŸŽ¨ /brand", "/brand "),
                ("ðŸ“§ /outreach", "/outreach "),
                ("ðŸš€ /campaign", "/campaign "),
                ("ðŸ–¼ï¸ /design", "/design "),
                ("âœï¸ /content", "/content "),
            ]
            for i, (label, cmd) in enumerate(ai_cmds):
                with cmd_cols[i % 5]:
                    if st.button(label, key=f"ai_qcmd_{i}", use_container_width=True):
                        st.session_state.prefill_command = cmd
                        st.session_state.show_command_palette = False
                        st.rerun()
            
            st.markdown("**ðŸ“š Knowledge Base**")
            kb_cols = st.columns(5)
            kb_cmds = [
                ("ðŸ“– /learn", "/learn "),
                ("ðŸ” /recall", "/recall "),
                ("ðŸ“Š /knowledge", "/knowledge"),
                ("ðŸ“ /status", "/status"),
                ("ðŸ—‘ï¸ /forget", "/forget"),
            ]
            for i, (label, cmd) in enumerate(kb_cmds):
                with kb_cols[i % 5]:
                    if st.button(label, key=f"kb_qcmd_{i}", use_container_width=True):
                        st.session_state.prefill_command = cmd
                        st.session_state.show_command_palette = False
                        st.rerun()
            
            st.markdown("**ðŸŽ¨ Media & Files**")
            cmd_cols2 = st.columns(4)
            quick_cmds = [
                ("ðŸŽ¨ /image", "/image "),
                ("ðŸŽ¬ /video", "/video "),
                ("ðŸŽµ /music", "/music "),
                ("ðŸ /python", "/python "),
                ("ðŸ“„ /pdf", "/pdf "),
                ("ðŸ“Š /json", "/json "),
                ("ðŸ“‘ /xlsx", "/xlsx "),
                ("ðŸ”— /chain", "/chain "),
            ]
            for i, (label, cmd) in enumerate(quick_cmds):
                with cmd_cols2[i % 4]:
                    if st.button(label, key=f"qcmd_{i}", use_container_width=True):
                        st.session_state.prefill_command = cmd
                        st.session_state.show_command_palette = False
                        st.rerun()
    
    # Input with prefill support
    prefill = st.session_state.pop('prefill_command', None)
    user_input = st.chat_input(
        placeholder="Message Otto... (type / for commands)",
        key="otto_chat_input"
    )
    
    if user_input:
        # Add user message
        st.session_state.otto_messages.append({
            'role': 'user',
            'content': user_input
        })
        
        # Check if this is a slash command
        slash_processor = get_slash_processor(replicate_api)
        if slash_processor.is_slash_command(user_input):
            # Execute slash command directly
            with st.spinner("âš¡ Executing command..."):
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                result = loop.run_until_complete(slash_processor.execute(user_input))
            
            # Format response
            if result.get('success'):
                artifacts = result.get('artifacts', [])
                
                # Build response message based on result type
                if result.get('type') == 'file':
                    response_msg = f"âœ… Created file: **{result.get('filename')}**\n\n"
                    if result.get('content_preview'):
                        response_msg += f"```\n{result.get('content_preview')}\n```"
                elif result.get('type') == 'media':
                    response_msg = f"âœ… Generated {result.get('media_type', 'media')}"
                elif result.get('type') == 'chain':
                    response_msg = result.get('message', 'âœ… Chain completed')
                else:
                    response_msg = result.get('message', 'âœ… Command completed successfully')
                
                st.session_state.otto_messages.append({
                    'role': 'assistant',
                    'content': response_msg,
                    'artifacts': artifacts
                })
            else:
                st.session_state.otto_messages.append({
                    'role': 'assistant',
                    'content': f"âŒ Error: {result.get('error', 'Unknown error')}",
                    'artifacts': []
                })
            
            st.rerun()
        
        # Otherwise, create and execute plan
        with st.spinner("ðŸ§  Otto is thinking..."):
            plan = engine.create_plan(user_input)
            st.session_state.otto_current_plan = plan
        
        # Show plan
        st.info(f"**Plan:** {len(plan.steps)} steps to complete")
        
        # Execute with progress
        progress_container = st.empty()
        results_container = st.container()
        
        async def execute_with_display():
            def progress_callback(update):
                with progress_container:
                    status_emoji = {"running": "â³", "completed": "âœ…", "failed": "âŒ"}.get(update['status'], "âšª")
                    st.markdown(f"""
                    <div class="otto-step otto-step-{update['status']}">
                        {status_emoji} Step {update['step']}/{update['total']}: {update['name']}
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Show result artifacts immediately with proper rendering
                    if update.get('result') and update['result'].get('artifacts'):
                        with results_container:
                            for artifact in update['result']['artifacts']:
                                artifact_type = artifact.get('type', '')
                                artifact_url = artifact.get('url', '')
                                
                                if artifact_type == 'image':
                                    st.image(artifact_url, caption="Generated Image", use_container_width=True)
                                elif artifact_type == 'video':
                                    st.video(artifact_url)
                                elif artifact_type == 'youtube_video':
                                    video_id = artifact.get('id', '')
                                    if video_id:
                                        st.markdown(f"""
                                        <iframe width="100%" height="400" 
                                            src="https://www.youtube.com/embed/{video_id}" 
                                            frameborder="0" allowfullscreen>
                                        </iframe>
                                        """, unsafe_allow_html=True)
                                    elif artifact_url:
                                        st.video(artifact_url)
                                elif artifact_type == 'printify_upload':
                                    st.success(f"âœ… Uploaded to Printify - Image ID: {artifact.get('id')}")
                                elif artifact_type in ['shopify_product', 'shopify_blog']:
                                    label = "Blog Post" if artifact_type == 'shopify_blog' else "Product"
                                    st.success(f"âœ… Published {label} to Shopify - ID: {artifact.get('id')}")
            
            return await engine.execute_plan(plan, progress_callback)
        
        # Run async execution
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        completed_plan = loop.run_until_complete(execute_with_display())
        
        # Collect all artifacts
        all_artifacts = []
        for step in completed_plan.steps:
            all_artifacts.extend(step.artifacts)
        
        # Add assistant response
        st.session_state.otto_messages.append({
            'role': 'assistant',
            'content': completed_plan.final_result,
            'artifacts': all_artifacts
        })
        
        st.rerun()
